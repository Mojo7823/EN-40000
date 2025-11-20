"""Conformance Claim section rendering helpers."""
from typing import Any, Iterable, Optional

from docx import Document
from docx.shared import Pt

from .html_converter import append_html_to_document


def append_conformance_claim_section(document: Document, conformance_data: Any):
  """Append Section 3 (Conformance Claim) to the document."""
  document.add_page_break()

  heading = document.add_paragraph()
  run = heading.add_run("3. Conformance Claim")
  run.font.size = Pt(20)
  run.font.bold = True
  heading.space_after = Pt(8)

  _add_standards_subsection(document, _get_nested(conformance_data, "standards_conformance"))
  _add_regulatory_subsection(document, _get_nested(conformance_data, "regulatory_conformance_html"))
  _add_conformance_level_subsection(document, _get_nested(conformance_data, "conformance_level_html"))


def _add_standards_subsection(document: Document, section_data: Any):
  sub_heading = document.add_paragraph()
  run = sub_heading.add_run("3.1 Standards Conformance")
  run.font.size = Pt(18)
  run.font.bold = True
  sub_heading.space_after = Pt(4)

  intro = document.add_paragraph("This product claims conformance to the following standards:")
  intro.space_after = Pt(6)

  primary_entry = _normalize_standard(_get_nested(section_data, "primary_standard"))
  _render_primary_standard(document, primary_entry)

  related_entries = _normalize_related_standards(_get_nested(section_data, "related_standards"))
  _render_related_standards(document, related_entries)

  include_other = bool(_get_nested(section_data, "include_other"))
  other_notes = str(_get_nested(section_data, "other_notes") or "").strip()
  _render_other_checkbox(document, include_other, other_notes)


def _render_primary_standard(document: Document, entry: Optional[dict]):
  label = document.add_paragraph()
  run = label.add_run("Primary Standard:")
  run.font.bold = True
  label.space_after = Pt(2)

  if entry:
    paragraph = document.add_paragraph(_format_standard(entry))
    paragraph.space_after = Pt(4)
    return

  placeholder = document.add_paragraph(
    "[Specify the primary standard that anchors this conformance claim.]"
  )
  placeholder.runs[0].font.italic = True
  placeholder.space_after = Pt(4)


def _render_related_standards(document: Document, entries: Iterable[dict]):
  heading = document.add_paragraph("Related Standards Applied:")
  heading.runs[0].font.bold = True
  heading.space_before = Pt(4)
  heading.space_after = Pt(2)

  entries = list(entries)
  if entries:
    for entry in entries:
      paragraph = document.add_paragraph(_format_standard(entry))
      paragraph.paragraph_format.left_indent = Pt(12)
      paragraph.space_after = Pt(2)
    return

  placeholder = document.add_paragraph("[List supplementary standards, profiles, or references applied.]")
  placeholder.paragraph_format.left_indent = Pt(12)
  placeholder.runs[0].font.italic = True
  placeholder.space_after = Pt(4)


def _render_other_checkbox(document: Document, include_other: bool, notes: str):
  symbol = "☑" if include_other and notes else "☐"
  line = document.add_paragraph()
  line.paragraph_format.left_indent = Pt(6)
  prefix = line.add_run(f"{symbol} Other: ")
  prefix.font.bold = True
  if notes:
    line.add_run(notes)
    return
  placeholder = line.add_run("[List any other standards applied]")
  placeholder.font.italic = True


def _add_regulatory_subsection(document: Document, html_content: Optional[str]):
  heading = document.add_paragraph()
  run = heading.add_run("3.2 Regulatory Conformance")
  run.font.size = Pt(18)
  run.font.bold = True
  heading.space_before = Pt(10)
  heading.space_after = Pt(4)

  if html_content:
    append_html_to_document(document, html_content)
    return

  placeholder = document.add_paragraph(
    "[Document the CRA clauses, EU directives, and industry-specific regulations satisfied by this product.]"
  )
  placeholder.runs[0].font.italic = True


def _add_conformance_level_subsection(document: Document, html_content: Optional[str]):
  heading = document.add_paragraph()
  run = heading.add_run("3.3 Conformance Level")
  run.font.size = Pt(18)
  run.font.bold = True
  heading.space_before = Pt(10)
  heading.space_after = Pt(4)

  if html_content:
    append_html_to_document(document, html_content)
    return

  placeholder = document.add_paragraph(
    "[State the claimed conformance level (e.g., Basic, Substantial, High) and summarize supporting evidence.]"
  )
  placeholder.runs[0].font.italic = True


def _format_standard(entry: dict) -> str:
  code = entry.get("code", "").strip()
  description = entry.get("description", "").strip()
  if code and description:
    return f"{code}: {description}"
  return code or description or "—"


def _normalize_standard(entry: Any) -> Optional[dict]:
  if not entry:
    return None
  if isinstance(entry, dict):
    data = entry
  else:
    data = {
      "code": getattr(entry, "code", None),
      "description": getattr(entry, "description", None),
    }
  code = (data.get("code") or "").strip()
  description = (data.get("description") or "").strip()
  if not code and not description:
    return None
  return {"code": code, "description": description}


def _normalize_related_standards(entries: Any) -> Iterable[dict]:
  if not entries:
    return []
  normalized = []
  for entry in entries:
    normalized_entry = _normalize_standard(entry)
    if normalized_entry:
      normalized.append(normalized_entry)
  return normalized


def _get_nested(source: Any, key: str):
  if not source:
    return None
  if isinstance(source, dict):
    return source.get(key)
  return getattr(source, key, None)

