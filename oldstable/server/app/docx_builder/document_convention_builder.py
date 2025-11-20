"""Document Convention section builder for DOCX generation."""
from pathlib import Path
from typing import List, Optional, Sequence

from docx import Document
from docx.shared import Inches, Pt

from .html_converter import append_html_to_document
from .section_builders import create_base_document

DEFAULT_TERMINOLOGY_ENTRIES = [
    {
        "term": "Product with digital elements",
        "definition": "A product or system containing digital components capable of processing, transmitting, or storing data.",
        "reference": "Regulation (EU) 2024/2847",
    },
    {
        "term": "Cybersecurity",
        "definition": "Ability to protect devices, networks, and services from attacks compromising confidentiality, integrity, or availability.",
        "reference": "prEN 40000-1-1",
    },
    {
        "term": "Vulnerability",
        "definition": "Weakness in design, implementation, or operation that can be exploited to compromise the product.",
        "reference": "Regulation (EU) 2024/2847",
    },
    {
        "term": "Risk",
        "definition": "Combination of the likelihood of a cybersecurity event occurring and the impact of that event.",
        "reference": "prEN 40000-1-1",
    },
    {
        "term": "RDPS",
        "definition": "Remote Digital Product Support services used to monitor, configure, or update the product post-deployment.",
        "reference": "EN 40000-1-2",
    },
    {
        "term": "SBOM",
        "definition": "Software Bill of Materials describing all software components within the product.",
        "reference": "ENISA SBOM requirements",
    },
    {
        "term": "IPRFU",
        "definition": "Installed Product Release and Field Updates covering deployed versions and maintenance activities.",
        "reference": "Internal engineering procedure",
    },
]


def build_document_convention_section(
    terminology_entries: Optional[Sequence[dict]] = None,
    evidence_notation_html: Optional[str] = None,
    requirement_notation_html: Optional[str] = None,
    assessment_verdicts_html: Optional[str] = None,
    output_dir: Optional[Path] = None,
    user_id: str = "default",
) -> Path:
    """Build the Document Convention section as a standalone DOCX file."""
    document = create_base_document()
    render_document_convention_to_document(
        document,
        terminology_entries=terminology_entries,
        evidence_notation_html=evidence_notation_html,
        requirement_notation_html=requirement_notation_html,
        assessment_verdicts_html=assessment_verdicts_html,
        start_on_new_page=False,
    )

    if output_dir is None:
        output_dir = Path("output")
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"document_convention_{user_id}.docx"
    document.save(str(output_path))
    return output_path


def render_document_convention_to_document(
    document: Document,
    *,
    terminology_entries: Optional[Sequence[dict]] = None,
    evidence_notation_html: Optional[str] = None,
    requirement_notation_html: Optional[str] = None,
    assessment_verdicts_html: Optional[str] = None,
    start_on_new_page: bool = False,
) -> None:
    """Append Section 4 (Document Conventions) to an existing Document."""
    if start_on_new_page:
        document.add_page_break()
    
    heading = document.add_paragraph()
    heading_run = heading.add_run("4. Document Conventions")
    heading_run.font.size = Pt(20)
    heading_run.font.bold = True
    heading.space_after = Pt(8)

    _render_terminology_section(document, terminology_entries)
    _render_evidence_notation_section(document, evidence_notation_html)
    _render_requirement_notation_section(document, requirement_notation_html)
    _render_assessment_verdicts_section(document, assessment_verdicts_html)


def _render_terminology_section(document: Document, entries: Optional[Sequence[dict]]) -> None:
    subheading = document.add_paragraph()
    subheading_run = subheading.add_run("4.1 Terminology")
    subheading_run.font.size = Pt(18)
    subheading_run.font.bold = True
    subheading.space_before = Pt(6)
    subheading.space_after = Pt(6)

    intro = document.add_paragraph(
        "All terms and definitions used in this report are consistent with:"
    )
    intro.paragraph_format.space_after = Pt(6)

    for reference in ("prEN 40000-1-1 (Vocabulary)", "Regulation (EU) 2024/2847"):
        paragraph = document.add_paragraph(reference, style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.space_after = Pt(3)

    document.add_paragraph()

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for label, cell in zip(["Term", "Definition", "Reference"], header_cells):
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(label)
        run.font.bold = True

    normalized_entries = _normalize_terminology_entries(entries)
    for entry in normalized_entries:
        row_cells = table.add_row().cells
        row_cells[0].text = entry["term"]
        row_cells[1].text = entry["definition"]
        row_cells[2].text = entry["reference"]

    document.add_paragraph()


def _render_evidence_notation_section(document: Document, html: Optional[str]) -> None:
    heading = document.add_paragraph()
    heading_run = heading.add_run("4.2 Evidence Notation")
    heading_run.font.size = Pt(18)
    heading_run.font.bold = True
    heading.space_before = Pt(6)
    heading.space_after = Pt(6)

    if html:
        append_html_to_document(document, html)


def _render_requirement_notation_section(document: Document, html: Optional[str]) -> None:
    heading = document.add_paragraph()
    heading_run = heading.add_run("4.3 Requirement Notation")
    heading_run.font.size = Pt(18)
    heading_run.font.bold = True
    heading.space_before = Pt(6)
    heading.space_after = Pt(6)

    if html:
        append_html_to_document(document, html)


def _render_assessment_verdicts_section(document: Document, html: Optional[str]) -> None:
    heading = document.add_paragraph()
    heading_run = heading.add_run("4.4 Assessment Verdicts")
    heading_run.font.size = Pt(18)
    heading_run.font.bold = True
    heading.space_before = Pt(6)
    heading.space_after = Pt(6)

    if html:
        append_html_to_document(document, html)


def _normalize_terminology_entries(entries: Optional[Sequence[dict]]) -> List[dict]:
    normalized: List[dict] = []
    for entry in entries or []:
        term = _safe_get(entry, "term")
        definition = _safe_get(entry, "definition")
        reference = _safe_get(entry, "reference")
        if term or definition or reference:
            normalized.append(
                {
                    "term": term or "—",
                    "definition": definition or "—",
                    "reference": reference or "—",
                }
            )
    if normalized:
        return normalized
    return [dict(item) for item in DEFAULT_TERMINOLOGY_ENTRIES]


def _safe_get(entry: object, key: str) -> str:
    if isinstance(entry, dict):
        value = entry.get(key)
    else:
        value = getattr(entry, key, None)
    if value is None:
        return ""
    return str(value).strip()
