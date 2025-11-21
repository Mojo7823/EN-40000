"""Rendering helpers for introduction sections (1.1–1.4)."""
from datetime import datetime
from typing import Any, Iterable

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from .html_converter import append_html_to_document

PURPOSE_SCOPE_PHASES = [
    "Concept and planning",
    "Design and development",
    "Implementation",
    "Verification and validation",
    "Production and distribution",
    "Deployment and operation",
    "Maintenance and support",
    "Decommissioning",
]

PRODUCT_DESCRIPTION_PLACEHOLDER = (
    "[Summarize the product’s architecture, deployment model, critical components, and safeguards.]"
)
KEY_FUNCTIONS_PLACEHOLDER = "[List the primary security or operational functions evaluated in this document.]"
TARGET_MARKET_PLACEHOLDER = "[e.g., Consumer, Enterprise, Industrial, Healthcare]"


class IntroductionSectionsRenderer:
    """Handles rendering of sections 1.1 through 1.4."""

    def __init__(self, document: Document):
        self.document = document

    def render(
        self,
        introduction_payload: Any,
        purpose_scope_payload: Any,
        product_identification_payload: Any,
        manufacturer_information_payload: Any,
    ):
        self.document.add_page_break()
        self._add_intro_heading()
        self._add_document_information_section(introduction_payload)
        self._add_purpose_scope_section(introduction_payload, purpose_scope_payload)
        self._add_product_identification_section(introduction_payload, product_identification_payload)
        self._add_manufacturer_information_section(manufacturer_information_payload)

    def _add_intro_heading(self):
        heading = self.document.add_paragraph()
        heading_run = heading.add_run("1. Introduction")
        heading_run.font.size = Pt(20)
        heading_run.font.bold = True
        heading.space_after = Pt(8)

        intro_para = self.document.add_paragraph(
            "This section captures the key Document Information required by the CRA Documentation toolkit."
        )
        intro_para.space_after = Pt(12)

    def _add_document_information_section(self, introduction_data: Any):
        heading = self.document.add_paragraph()
        heading_run = heading.add_run("1.1 Document Information")
        heading_run.font.size = Pt(18)
        heading_run.font.bold = True
        heading.space_before = Pt(6)
        heading.space_after = Pt(8)

        fields = [
            ("Product Name", "product_name"),
            ("Product Version", "product_version"),
            ("Product Type / Category", "product_type"),
            ("Manufacturer", "manufacturer"),
            ("Manufacturer Address", "manufacturer_address"),
            ("Status", "status"),
            ("Prepared By", "prepared_by"),
            ("Reviewed By", "reviewed_by"),
            ("Approved By", "approved_by"),
        ]

        table = self.document.add_table(rows=len(fields), cols=2)
        table.style = "Table Grid"
        table.autofit = True

        for idx, (label, key) in enumerate(fields):
            label_cell = table.cell(idx, 0)
            value_cell = table.cell(idx, 1)

            _write_table_cell(label_cell, label, bold=True)
            value = _get_intro_value(introduction_data, key)
            _write_table_cell(value_cell, value or "—")

        self.document.add_paragraph().space_after = Pt(6)

    def _add_purpose_scope_section(self, introduction_data: Any, purpose_data: Any):
        self.document.add_page_break()

        heading = self.document.add_paragraph()
        heading_run = heading.add_run("1.2 Purpose and Scope")
        heading_run.font.size = Pt(18)
        heading_run.font.bold = True
        heading.space_after = Pt(8)

        reference = self.document.add_paragraph("[Reference: Clause 1 - Scope]")
        reference.runs[0].font.bold = True
        reference.space_after = Pt(10)

        product_name = _get_intro_value(introduction_data, "product_name")
        if not product_name:
            product_name = _get_purpose_value(purpose_data, "product_name") or "the Product"

        intro_paragraph = self.document.add_paragraph(
            (
                "This conformance report documents the assessment of "
                f"{product_name} against the requirements specified in EN 40000-1-2-2025: "
                "‘Cybersecurity requirements for products with digital elements – Part 1-2: Principles for cyber resilience.’"
            )
        )
        intro_paragraph.space_after = Pt(8)

        follow_up = self.document.add_paragraph(
            (
                "The standard specifies general cybersecurity principles and risk management activities applicable to all products "
                "with digital elements throughout their lifecycle. This report demonstrates how "
                f"{product_name} ensures and maintains an appropriate level of cybersecurity based on identified risks."
            )
        )
        follow_up.space_after = Pt(12)

        scope_heading = self.document.add_paragraph("Scope of Assessment:")
        scope_heading.runs[0].font.bold = True
        scope_heading.space_after = Pt(4)

        scope_intro = self.document.add_paragraph(
            f"This assessment covers the following lifecycle phases of {product_name}:"
        )
        scope_intro.space_after = Pt(6)

        selections = set(_get_purpose_value(purpose_data, "scope_selections", []))
        for phase in PURPOSE_SCOPE_PHASES:
            paragraph = self.document.add_paragraph()
            symbol = "☑" if phase in selections else "☐"
            paragraph.add_run(f"{symbol} {phase}")
            paragraph.space_after = Pt(2)

        assessment_heading = self.document.add_paragraph("Assessment Period:")
        assessment_heading.runs[0].font.bold = True
        assessment_heading.space_before = Pt(10)
        assessment_heading.space_after = Pt(2)

        period_text = _format_assessment_period(
            _get_purpose_value(purpose_data, "assessment_start"),
            _get_purpose_value(purpose_data, "assessment_end"),
        )
        self.document.add_paragraph(period_text or "[Start Date] to [End Date]")

        methodology_heading = self.document.add_paragraph("Assessment Methodology:")
        methodology_heading.runs[0].font.bold = True
        methodology_heading.space_before = Pt(10)
        methodology_heading.space_after = Pt(4)

        methodology_html = _get_purpose_value(purpose_data, "methodology_html")
        if methodology_html:
            append_html_to_document(self.document, methodology_html)
        else:
            placeholder = self.document.add_paragraph(
                "[Describe the methodology used – e.g., document review, interviews, technical testing, code review, penetration testing]."
            )
            placeholder.runs[0].font.italic = True

    def _add_product_identification_section(self, introduction_data: Any, product_data: Any):
        self.document.add_page_break()

        heading = self.document.add_paragraph()
        heading_run = heading.add_run("1.3 Product Identification")
        heading_run.font.size = Pt(18)
        heading_run.font.bold = True
        heading.space_after = Pt(8)

        fields = [
            ("Product Name", _get_intro_value(introduction_data, "product_name")),
            ("Product Version", _get_intro_value(introduction_data, "product_version")),
            ("Product Type / Category", _get_intro_value(introduction_data, "product_type")),
        ]

        table = self.document.add_table(rows=len(fields), cols=2)
        table.style = "Table Grid"
        table.autofit = True

        for idx, (label, value) in enumerate(fields):
            label_cell = table.cell(idx, 0)
            value_cell = table.cell(idx, 1)

            _write_table_cell(label_cell, label, bold=True)
            _write_table_cell(value_cell, value or "—")

        self.document.add_paragraph().space_after = Pt(4)

        description_heading = self.document.add_paragraph("Product Description:")
        description_heading.runs[0].font.bold = True
        description_heading.space_before = Pt(8)
        description_heading.space_after = Pt(4)

        description_html = _get_product_ident_value(product_data, "product_description_html")
        if description_html:
            append_html_to_document(self.document, description_html)
        else:
            placeholder = self.document.add_paragraph(PRODUCT_DESCRIPTION_PLACEHOLDER)
            placeholder.runs[0].font.italic = True

        functions_heading = self.document.add_paragraph("Key Product Functions:")
        functions_heading.runs[0].font.bold = True
        functions_heading.space_before = Pt(10)
        functions_heading.space_after = Pt(4)

        functions_html = _get_product_ident_value(product_data, "key_functions_html")
        if functions_html:
            append_html_to_document(self.document, functions_html)
        else:
            placeholder = self.document.add_paragraph(KEY_FUNCTIONS_PLACEHOLDER)
            placeholder.runs[0].font.italic = True

        target_heading = self.document.add_paragraph("Target Market:")
        target_heading.runs[0].font.bold = True
        target_heading.space_before = Pt(10)
        target_heading.space_after = Pt(2)

        target_value = _get_product_ident_value(product_data, "target_market")
        if target_value:
            for line in _split_lines(target_value):
                line_para = self.document.add_paragraph(line)
                line_para.paragraph_format.space_after = Pt(2)
        else:
            placeholder = self.document.add_paragraph(TARGET_MARKET_PLACEHOLDER)
            placeholder.runs[0].font.italic = True

    def _add_manufacturer_information_section(self, manufacturer_data: Any):
        self.document.add_page_break()

        heading = self.document.add_paragraph()
        heading_run = heading.add_run("1.4 Manufacturer Information")
        heading_run.font.size = Pt(18)
        heading_run.font.bold = True
        heading.space_after = Pt(8)

        fields = [
            ("Legal Entity", _get_manufacturer_value(manufacturer_data, "legal_entity")),
            ("Registration Number", _get_manufacturer_value(manufacturer_data, "registration_number")),
            ("Address", _get_manufacturer_value(manufacturer_data, "address")),
            ("Contact Person", _get_manufacturer_value(manufacturer_data, "contact_person")),
            ("Phone", _get_manufacturer_value(manufacturer_data, "phone")),
        ]

        table = self.document.add_table(rows=len(fields), cols=2)
        table.style = "Table Grid"
        table.autofit = True

        shaded_rows = {1, 3}
        for idx, (label, value) in enumerate(fields):
            label_cell = table.cell(idx, 0)
            value_cell = table.cell(idx, 1)

            _write_table_cell(label_cell, label, bold=True)
            _write_table_cell(value_cell, value or "—")

            if idx in shaded_rows:
                _apply_cell_shading(label_cell, "f3f4f6")
                _apply_cell_shading(value_cell, "f3f4f6")

        self.document.add_paragraph().space_after = Pt(6)


def _split_lines(value: Any) -> Iterable[str]:
    if not value:
        return []
    return [line.strip() for line in str(value).splitlines() if line.strip()]


def _write_table_cell(cell, text: str, *, bold: bool = False):
    paragraph = cell.paragraphs[0]
    paragraph.text = ""
    run = paragraph.add_run(text)
    run.font.bold = bold


def _get_intro_value(source: Any, key: str) -> str:
    if not source:
        return ""
    if isinstance(source, dict):
        value = source.get(key)
    else:
        value = getattr(source, key, None)
    if value is None:
        return ""
    return str(value).strip()


def _get_purpose_value(source: Any, key: str, default=None):
    if not source:
        return default
    if isinstance(source, dict):
        value = source.get(key, default)
    else:
        value = getattr(source, key, default)
    return default if value is None else value


def _get_product_ident_value(source: Any, key: str) -> str:
    if not source:
        return ""
    if isinstance(source, dict):
        value = source.get(key)
    else:
        value = getattr(source, key, None)
    if value is None:
        return ""
    return str(value).strip()


def _get_manufacturer_value(source: Any, key: str) -> str:
    if not source:
        return ""
    if isinstance(source, dict):
        value = source.get(key)
    else:
        value = getattr(source, key, None)
    if value is None:
        return ""
    return str(value).strip()


def _format_assessment_period(start: str = None, end: str = None) -> str:
    start_label = _format_period_date(start)
    end_label = _format_period_date(end)
    if not start_label and not end_label:
        return ""
    if not start_label:
        return f"Through {end_label}"
    if not end_label:
        return f"{start_label} onward"
    return f"{start_label} to {end_label}"


def _format_period_date(value: str = None) -> str:
    if not value:
        return ""
    try:
        parsed = datetime.fromisoformat(value)
    except Exception:
        return value
    return parsed.strftime("%B %d, %Y")


def _apply_cell_shading(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)
