"""Risk Management Elements section builder."""
from typing import Iterable, Optional
import re

from docx import Document
from docx.shared import Pt

from .html_converter import append_html_to_document

EVIDENCE_STATUS_LABELS = {
    "complete": "Complete",
    "in_progress": "In Progress",
    "not_started": "Not Started",
}

ASSESSMENT_STATUS_LABELS = {
    "pass": "Pass",
    "fail": "Fail",
    "na": "N/A",
    "not_assessed": "Not Assessed",
}

OVERALL_VERDICT_LABELS = {
    "pass": "PASS",
    "partial": "PARTIAL",
    "fail": "FAIL",
    "na": "N/A",
    "not_assessed": "Not Assessed",
}

NC_SEVERITY_LABELS = {
    "minor": "Minor",
    "major": "Major",
    "critical": "Critical",
}

# Product Context Assessment Requirements (Clause 6.2)
PRODUCT_CONTEXT_REQUIREMENTS = [
    # 6.2.2 Input Documentation
    {"id": "6.2.2-a", "subsection": "6.2.2", "label": "Input Documentation", "description": "Functional use cases (user stories) documented"},
    {"id": "6.2.2-b", "subsection": "6.2.2", "label": "Input Documentation", "description": "User types identified and documented"},
    {"id": "6.2.2-c", "subsection": "6.2.2", "label": "Input Documentation", "description": "Operational environment documented"},
    {"id": "6.2.2-d", "subsection": "6.2.2", "label": "Input Documentation", "description": "Product architecture documented"},
    {"id": "6.2.2-e", "subsection": "6.2.2", "label": "Input Documentation", "description": "Third-party components identified"},
    # 6.2.3 Product Context Requirements
    {"id": "6.2.3-a", "subsection": "6.2.3", "label": "Product Context Requirements", "description": "Intended purpose and reasonably foreseeable use (IPRFU) defined"},
    {"id": "6.2.3-b", "subsection": "6.2.3", "label": "Product Context Requirements", "description": "Product functions documented"},
    {"id": "6.2.3-c", "subsection": "6.2.3", "label": "Product Context Requirements", "description": "Operational environment specified"},
    {"id": "6.2.3-d", "subsection": "6.2.3", "label": "Product Context Requirements", "description": "User types and characteristics defined"},
    {"id": "6.2.3-e", "subsection": "6.2.3", "label": "Product Context Requirements", "description": "Product architecture described"},
    {"id": "6.2.3-f", "subsection": "6.2.3", "label": "Product Context Requirements", "description": "RDPS dependencies recorded (if applicable)"},
    {"id": "6.2.3-g", "subsection": "6.2.3", "label": "Product Context Requirements", "description": "Data processing activities identified (if applicable)"},
    {"id": "6.2.3-h", "subsection": "6.2.3", "label": "Product Context Requirements", "description": "External interfaces documented (if applicable)"},
    {"id": "6.2.3-i", "subsection": "6.2.3", "label": "Product Context Requirements", "description": "Security assumptions documented"},
    # 6.2.4 Output
    {"id": "6.2.4-a", "subsection": "6.2.4", "label": "Output", "description": "Product context documentation complete"},
    {"id": "6.2.4-b", "subsection": "6.2.4", "label": "Output", "description": "Product context review record available"},
    # 6.2.5 Assessment Criteria
    {"id": "6.2.5-a", "subsection": "6.2.5", "label": "Assessment Criteria", "description": "Product context is complete and consistent"},
    {"id": "6.2.5-b", "subsection": "6.2.5", "label": "Assessment Criteria", "description": "Product context is traceable to requirements"},
]


def append_risk_management_section(
    document: Document, payload: Optional[object], product_name: str = "[Product Name]"
) -> None:
    """Append Section 5 - Risk Management Elements to the document."""
    if not payload:
        return

    general_html = _extract_value(payload, "general_approach_html")
    product_context_payload = getattr(payload, "product_context", None)
    product_function_payload = getattr(payload, "product_function", None)
    operational_environment_payload = getattr(payload, "operational_environment", None)
    product_architecture_payload = getattr(payload, "product_architecture", None)
    product_user_description_payload = getattr(payload, "product_user_description", None)
    product_context_assessment_payload = getattr(payload, "product_context_assessment", None)
    has_product_context = _product_context_has_content(product_context_payload)
    has_product_function = _product_function_has_content(product_function_payload)
    has_operational_environment = _operational_environment_has_content(operational_environment_payload)
    has_product_architecture = _product_architecture_has_content(product_architecture_payload)
    has_product_user_description = _product_user_description_has_content(product_user_description_payload)
    has_product_context_assessment = _product_context_assessment_has_content(product_context_assessment_payload)

    if not general_html and not has_product_context and not has_product_function and not has_operational_environment and not has_product_architecture and not has_product_user_description and not has_product_context_assessment:
        return

    document.add_page_break()

    # Section 5 heading
    heading = document.add_paragraph()
    heading_run = heading.add_run("5. RISK MANAGEMENT ELEMENTS")
    heading_run.font.size = Pt(20)
    heading_run.font.bold = True
    heading.space_after = Pt(8)

    # Main clause reference
    reference = document.add_paragraph("[Reference: Clause 6 - Risk management elements]")
    reference.runs[0].font.bold = True
    reference.space_after = Pt(10)

    if general_html:
        # 5.1 subheading
        subheading = document.add_paragraph()
        subheading_run = subheading.add_run("5.1 General Approach to Risk Management")
        subheading_run.font.size = Pt(18)
        subheading_run.font.bold = True
        subheading.space_before = Pt(6)
        subheading.space_after = Pt(6)

        # Clause reference for 5.1
        clause_reference = document.add_paragraph("[Reference: Clause 6.1 - General]")
        clause_reference.runs[0].font.bold = True
        clause_reference.space_after = Pt(10)

        # Intro paragraph with product name
        intro_text = (
            f"This section describes how {product_name} applies risk management throughout its lifecycle "
            "to ensure an appropriate level of cybersecurity."
        )
        intro_para = document.add_paragraph(intro_text)
        intro_para.space_after = Pt(10)

        # Risk Management Framework Applied label
        framework_label = document.add_paragraph()
        framework_run = framework_label.add_run("Risk Management Framework Applied:")
        framework_run.font.bold = True
        framework_label.space_after = Pt(6)

        # WYSIWYG content
        append_html_to_document(document, general_html)

    if has_product_context:
        _append_product_context_section(document, product_context_payload, product_name)

    if has_product_function:
        _append_product_function_section(document, product_function_payload)

    if has_operational_environment:
        _append_operational_environment_section(document, operational_environment_payload)

    if has_product_architecture:
        _append_product_architecture_section(document, product_architecture_payload)

    if has_product_user_description:
        _append_product_user_description_section(document, product_user_description_payload)

    if has_product_context_assessment:
        _append_product_context_assessment_section(document, product_context_assessment_payload)


def _append_product_context_section(
    document: Document, payload: Optional[object], product_name: str = "[Product Name]"
) -> None:
    from docx.shared import RGBColor
    
    # Start 5.2 Product Context on a new page
    document.add_page_break()
    
    heading = document.add_paragraph()
    heading_run = heading.add_run("5.2 Product Context")
    heading_run.font.size = Pt(18)
    heading_run.font.bold = True
    heading.space_after = Pt(6)

    reference = document.add_paragraph("[Reference: Clause 6.2 - Product Context]")
    reference.runs[0].font.bold = True
    reference.space_after = Pt(10)

    subsection = document.add_paragraph()
    subsection_run = subsection.add_run("5.2.1 Intended Purpose and Reasonably Foreseeable Use")
    subsection_run.font.size = Pt(16)
    subsection_run.font.bold = True
    subsection.space_before = Pt(6)
    subsection.space_after = Pt(4)

    clause_reference = document.add_paragraph(
        "[Reference: Clause 6.2.1.2 - Product intended purpose and reasonable foreseeable use]"
    )
    clause_reference.runs[0].font.bold = True
    clause_reference.space_after = Pt(8)

    # Blue requirement section
    req_para = document.add_paragraph()
    req_run = req_para.add_run("Requirement [Clause 6.2.3]:")
    req_run.font.color.rgb = RGBColor(0, 0, 255)
    req_para.space_after = Pt(4)
    
    req_text = document.add_paragraph(
        '"The product context shall be identified and recorded based on:'
    )
    req_text.runs[0].font.color.rgb = RGBColor(0, 0, 255)
    req_text.space_after = Pt(2)
    
    bullets = [
        "the product's IPRFU;",
        "the product's functions;",
        "the product's operational environment of use;",
        "the product's architecture overview;",
        "the product's user descriptions."
    ]
    for bullet in bullets:
        bullet_para = document.add_paragraph(bullet, style='List Bullet')
        bullet_para.runs[0].font.color.rgb = RGBColor(0, 0, 255)
        bullet_para.space_after = Pt(2)
    
    closing_quote = document.add_paragraph('"')
    closing_quote.runs[0].font.color.rgb = RGBColor(0, 0, 255)
    closing_quote.space_after = Pt(12)

    _append_product_context_block(
        document,
        "Intended Purpose:",
        _extract_value(payload, "intended_purpose_html"),
    )
    _append_product_context_block(
        document,
        "Reasonably Foreseeable Use & Misuse:",
        _extract_value(payload, "foreseeable_use_html"),
    )

    _append_evidence_tracker(document, getattr(payload, "evidence_entries", None))


def _append_product_context_block(document: Document, title: str, html_content: Optional[str]) -> None:
    if not html_content:
        return
    subheading = document.add_paragraph()
    subheading_run = subheading.add_run(title)
    subheading_run.font.size = Pt(13)
    subheading_run.font.bold = True
    subheading.space_before = Pt(6)
    subheading.space_after = Pt(4)
    append_html_to_document(document, html_content)


def _append_evidence_tracker(document: Document, payload: Optional[Iterable[object]]) -> None:
    entries = _normalize_evidence_entries(payload)
    if not entries:
        return

    heading = document.add_paragraph()
    heading_run = heading.add_run("Evidence Reference:")
    heading_run.font.size = Pt(13)
    heading_run.font.bold = True
    heading.space_before = Pt(10)
    heading.space_after = Pt(4)

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Evidence Reference", "Title / Artifact", "Status", "Notes"]
    header_cells = table.rows[0].cells
    for idx, label in enumerate(headers):
        paragraph = header_cells[idx].paragraphs[0]
        run = paragraph.add_run(label)
        run.font.bold = True

    for entry in entries:
        row = table.add_row().cells
        row[0].text = entry["reference"] or ""
        row[1].text = entry["title"] or ""
        row[2].text = entry["status_label"]
        row[3].text = entry["notes"] or ""


def _normalize_evidence_entries(payload: Optional[Iterable[object]]):
    if not payload:
        return []
    normalized = []
    for entry in payload:
        reference = _extract_value(entry, "reference_id")
        title = _extract_value(entry, "title")
        status = _extract_value(entry, "status") or "not_started"
        notes = _strip_html(_extract_value(entry, "notes_html"))
        if not any([reference, title, notes]):
            continue
        normalized.append(
            {
                "reference": reference,
                "title": title,
                "status_label": EVIDENCE_STATUS_LABELS.get(status, status.replace("_", " ").title()),
                "notes": notes,
            }
        )
    return normalized


def _product_context_has_content(payload: Optional[object]) -> bool:
    if not payload:
        return False
    return any(
        [
            bool(_extract_value(payload, "intended_purpose_html")),
            bool(_extract_value(payload, "specific_intended_uses_html")),
            bool(_extract_value(payload, "foreseeable_use_html")),
            bool(_normalize_evidence_entries(getattr(payload, "evidence_entries", None))),
        ]
    )


def _product_function_has_content(payload: Optional[object]) -> bool:
    if not payload:
        return False
    return any(
        [
            bool(_extract_value(payload, "primary_functions_html")),
            bool(_extract_value(payload, "security_functions_html")),
            bool(_normalize_evidence_entries(getattr(payload, "evidence_entries", None))),
        ]
    )


def _operational_environment_has_content(payload: Optional[object]) -> bool:
    if not payload:
        return False
    return any(
        [
            bool(_extract_value(payload, "physical_environment_html")),
            bool(_extract_value(payload, "network_environment_html")),
            bool(_extract_value(payload, "system_environment_html")),
            bool(_extract_value(payload, "operational_constraints_html")),
            bool(_extract_value(payload, "rdps_environment_html")),
            bool(_normalize_evidence_entries(getattr(payload, "evidence_entries", None))),
        ]
    )


def _product_architecture_has_content(payload: Optional[object]) -> bool:
    if not payload:
        return False
    return any(
        [
            bool(_extract_value(payload, "architecture_description_html")),
            bool(getattr(payload, "hardware_components", None)),
            bool(getattr(payload, "software_components", None)),
            bool(getattr(payload, "rdps_components", None)),
            bool(getattr(payload, "component_interfaces", None)),
            bool(_extract_value(payload, "architecture_diagram_html")),
            bool(_normalize_evidence_entries(getattr(payload, "evidence_entries", None))),
        ]
    )


def _product_user_description_has_content(payload: Optional[object]) -> bool:
    if not payload:
        return False
    return any(
        [
            bool(_extract_value(payload, "user_description_html")),
            getattr(payload, "no_rdps", False),
            bool(_extract_value(payload, "rdps_considerations_html")),
            bool(_normalize_evidence_entries(getattr(payload, "evidence_entries", None))),
        ]
    )


def _product_context_assessment_has_content(payload: Optional[object]) -> bool:
    """Check if Product Context Assessment section has content."""
    if not payload:
        return False
    
    # Check if any assessments have been filled in
    assessments = getattr(payload, "assessments", None) or []
    has_assessments = any(
        _extract_value(a, "status") not in (None, "not_assessed") or
        _extract_value(a, "evidence_id") or
        _extract_value(a, "comments_html")
        for a in assessments
    )
    
    # Check if verdict is set
    verdict = _extract_value(payload, "overall_verdict")
    has_verdict = verdict and verdict != "not_assessed"
    
    # Check if summary has content
    has_summary = bool(_extract_value(payload, "summary_of_findings_html"))
    
    # Check if there are non-conformities
    non_conformities = getattr(payload, "non_conformities", None) or []
    has_non_conformities = len(non_conformities) > 0
    
    return any([has_assessments, has_verdict, has_summary, has_non_conformities])


def _append_product_function_section(document: Document, payload: Optional[object]) -> None:
    if not payload:
        return

    # Start 5.2.2 Product Functions on a new page
    document.add_page_break()

    heading = document.add_paragraph()
    heading_run = heading.add_run("5.2.2 Product Functions")
    heading_run.font.size = Pt(16)
    heading_run.font.bold = True
    heading.space_after = Pt(6)

    reference = document.add_paragraph("[Reference: Clause 6.2.1.3 - Product functions]")
    reference.runs[0].font.bold = True
    reference.space_after = Pt(10)

    primary_html = _extract_value(payload, "primary_functions_html")
    if primary_html:
        append_html_to_document(document, primary_html)

    security_html = _extract_value(payload, "security_functions_html")
    if security_html:
        security_label = document.add_paragraph()
        security_run = security_label.add_run("Security Functions")
        security_run.font.size = Pt(13)
        security_run.font.bold = True
        security_label.space_before = Pt(10)
        security_label.space_after = Pt(6)
        append_html_to_document(document, security_html)

    _append_evidence_tracker(document, getattr(payload, "evidence_entries", None))


def _append_operational_environment_section(document: Document, payload: Optional[object]) -> None:
    from docx.shared import RGBColor
    
    if not payload:
        return

    # Start 5.2.3 Product Operational Environment on a new page
    document.add_page_break()

    heading = document.add_paragraph()
    heading_run = heading.add_run("5.2.3 Product Operational Environment")
    heading_run.font.size = Pt(16)
    heading_run.font.bold = True
    heading.space_after = Pt(6)

    reference = document.add_paragraph("[Reference: Clause 6.2.1.4 - Product operational environment]")
    reference.runs[0].font.bold = True
    reference.space_after = Pt(10)

    # Physical Environment
    physical_html = _extract_value(payload, "physical_environment_html")
    if physical_html:
        _append_environment_block(document, "Physical Environment:", physical_html)

    # Network Environment
    network_html = _extract_value(payload, "network_environment_html")
    if network_html:
        _append_environment_block(document, "Network Environment:", network_html)

    # System Environment
    system_html = _extract_value(payload, "system_environment_html")
    if system_html:
        _append_environment_block(document, "System Environment:", system_html)

    # Operational Constraints
    constraints_html = _extract_value(payload, "operational_constraints_html")
    if constraints_html:
        _append_environment_block(document, "Operational Constraints:", constraints_html)

    # RDPS Environment
    rdps_html = _extract_value(payload, "rdps_environment_html")
    if rdps_html:
        rdps_label = document.add_paragraph()
        rdps_run = rdps_label.add_run("RDPS Environment (if applicable):")
        rdps_run.font.size = Pt(13)
        rdps_run.font.bold = True
        rdps_label.space_before = Pt(10)
        rdps_label.space_after = Pt(4)

        # Blue requirement text
        req_para = document.add_paragraph()
        req_run = req_para.add_run('Requirement [Clause 6.2.3]: ')
        req_run.font.color.rgb = RGBColor(0, 0, 255)
        req_run.font.italic = True
        req_text = req_para.add_run(
            '"If applicable, RDPS dependencies (e.g. a concise dependency map identifying operator, '
            'data exchanged, trust/authentication, and defined degraded modes) shall be recorded."'
        )
        req_text.font.color.rgb = RGBColor(0, 0, 255)
        req_text.font.italic = True
        req_para.space_after = Pt(8)

        append_html_to_document(document, rdps_html)

    _append_evidence_tracker(document, getattr(payload, "evidence_entries", None))


def _append_product_architecture_section(document: Document, payload: Optional[object]) -> None:
    from docx.shared import RGBColor
    
    if not payload:
        return

    # Start 5.2.4 Product Architecture on a new page
    document.add_page_break()

    heading = document.add_paragraph()
    heading_run = heading.add_run("5.2.4 Product Architecture Overview")
    heading_run.font.size = Pt(16)
    heading_run.font.bold = True
    heading.space_after = Pt(6)

    reference = document.add_paragraph("[Reference: Clause 6.2.3 - Product architecture]")
    reference.runs[0].font.bold = True
    reference.space_after = Pt(10)

    # Architecture Description
    arch_desc_html = _extract_value(payload, "architecture_description_html")
    if arch_desc_html:
        desc_label = document.add_paragraph()
        desc_run = desc_label.add_run("Architecture Description:")
        desc_run.font.size = Pt(13)
        desc_run.font.bold = True
        desc_label.space_before = Pt(6)
        desc_label.space_after = Pt(4)
        append_html_to_document(document, arch_desc_html)

    # Hardware Components
    no_hardware = getattr(payload, "no_hardware_components", False)
    hardware_components = getattr(payload, "hardware_components", None) or []
    
    hw_label = document.add_paragraph()
    hw_run = hw_label.add_run("Hardware Components:")
    hw_run.font.size = Pt(13)
    hw_run.font.bold = True
    hw_label.space_before = Pt(10)
    hw_label.space_after = Pt(4)

    if no_hardware:
        no_hw_para = document.add_paragraph("This product does not have any hardware component.")
        no_hw_para.runs[0].font.italic = True
        no_hw_para.space_after = Pt(8)
    elif hardware_components:
        _append_hardware_table(document, hardware_components)
    else:
        no_hw_para = document.add_paragraph("No hardware components specified.")
        no_hw_para.runs[0].font.italic = True
        no_hw_para.space_after = Pt(8)

    # Software Components
    software_components = getattr(payload, "software_components", None) or []
    
    sw_label = document.add_paragraph()
    sw_run = sw_label.add_run("Software Components:")
    sw_run.font.size = Pt(13)
    sw_run.font.bold = True
    sw_label.space_before = Pt(10)
    sw_label.space_after = Pt(4)

    if software_components:
        _append_software_table(document, software_components)
    else:
        no_sw_para = document.add_paragraph("No software components specified.")
        no_sw_para.runs[0].font.italic = True
        no_sw_para.space_after = Pt(8)

    # RDPS Components
    no_rdps = getattr(payload, "no_rdps_components", False)
    rdps_components = getattr(payload, "rdps_components", None) or []
    
    rdps_label = document.add_paragraph()
    rdps_run = rdps_label.add_run("RDPS Components (if applicable):")
    rdps_run.font.size = Pt(13)
    rdps_run.font.bold = True
    rdps_label.space_before = Pt(10)
    rdps_label.space_after = Pt(4)

    # Blue requirement text for RDPS
    req_para = document.add_paragraph()
    req_run = req_para.add_run("Requirement [Clause 6.2.3]: ")
    req_run.font.color.rgb = RGBColor(0, 0, 255)
    req_run.font.italic = True
    req_text = req_para.add_run(
        '"Where a product function relies on an RDPS, the manufacturer shall include the RDPS '
        'in the product context determination. The product including its RDPS, third-party or not, '
        'shall be treated as a single system."'
    )
    req_text.font.color.rgb = RGBColor(0, 0, 255)
    req_text.font.italic = True
    req_para.space_after = Pt(8)

    if no_rdps:
        no_rdps_para = document.add_paragraph("This product does not rely on any RDPS.")
        no_rdps_para.runs[0].font.italic = True
        no_rdps_para.space_after = Pt(8)
    elif rdps_components:
        _append_rdps_table(document, rdps_components)
    else:
        no_rdps_para = document.add_paragraph("No RDPS components specified.")
        no_rdps_para.runs[0].font.italic = True
        no_rdps_para.space_after = Pt(8)

    # Component Interfaces
    component_interfaces = getattr(payload, "component_interfaces", None) or []
    
    iface_label = document.add_paragraph()
    iface_run = iface_label.add_run("Component Interfaces:")
    iface_run.font.size = Pt(13)
    iface_run.font.bold = True
    iface_label.space_before = Pt(10)
    iface_label.space_after = Pt(4)

    if component_interfaces:
        _append_interface_table(document, component_interfaces)
    else:
        no_iface_para = document.add_paragraph("No component interfaces specified.")
        no_iface_para.runs[0].font.italic = True
        no_iface_para.space_after = Pt(8)

    # Architecture Diagram
    arch_diagram_html = _extract_value(payload, "architecture_diagram_html")
    if arch_diagram_html:
        diagram_label = document.add_paragraph()
        diagram_run = diagram_label.add_run("Architecture Diagram:")
        diagram_run.font.size = Pt(13)
        diagram_run.font.bold = True
        diagram_label.space_before = Pt(10)
        diagram_label.space_after = Pt(4)
        append_html_to_document(document, arch_diagram_html)

    _append_evidence_tracker(document, getattr(payload, "evidence_entries", None))


def _append_hardware_table(document: Document, components: list) -> None:
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Component Name", "Function", "Interfaces", "Security Functions"]
    header_cells = table.rows[0].cells
    for idx, label in enumerate(headers):
        paragraph = header_cells[idx].paragraphs[0]
        run = paragraph.add_run(label)
        run.font.bold = True

    for comp in components:
        row = table.add_row().cells
        row[0].text = _extract_value(comp, "component_name") or ""
        row[1].text = _extract_value(comp, "function") or ""
        row[2].text = _extract_value(comp, "interfaces") or ""
        row[3].text = _extract_value(comp, "security_functions") or ""


def _append_software_table(document: Document, components: list) -> None:
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Type", "Function", "Third-Party", "Interfaces", "Security Functions"]
    header_cells = table.rows[0].cells
    for idx, label in enumerate(headers):
        paragraph = header_cells[idx].paragraphs[0]
        run = paragraph.add_run(label)
        run.font.bold = True

    for comp in components:
        row = table.add_row().cells
        row[0].text = _extract_value(comp, "type") or ""
        row[1].text = _extract_value(comp, "function") or ""
        third_party = getattr(comp, "third_party", None) if hasattr(comp, "third_party") else comp.get("third_party", False)
        row[2].text = "Yes" if third_party else "No"
        row[3].text = _extract_value(comp, "interfaces") or ""
        row[4].text = _extract_value(comp, "security_functions") or ""


def _append_rdps_table(document: Document, components: list) -> None:
    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["RDPS Component", "Provider", "Function", "Location", "Dev Responsibility", "Op Responsibility"]
    header_cells = table.rows[0].cells
    for idx, label in enumerate(headers):
        paragraph = header_cells[idx].paragraphs[0]
        run = paragraph.add_run(label)
        run.font.bold = True

    for comp in components:
        row = table.add_row().cells
        row[0].text = _extract_value(comp, "component") or ""
        row[1].text = _extract_value(comp, "provider") or ""
        row[2].text = _extract_value(comp, "function") or ""
        row[3].text = _extract_value(comp, "location") or ""
        row[4].text = _extract_value(comp, "development_responsibility") or ""
        row[5].text = _extract_value(comp, "operation_responsibility") or ""


def _append_interface_table(document: Document, interfaces: list) -> None:
    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Interface", "Component A", "Component B", "Protocol", "Authentication", "Data Exchanged"]
    header_cells = table.rows[0].cells
    for idx, label in enumerate(headers):
        paragraph = header_cells[idx].paragraphs[0]
        run = paragraph.add_run(label)
        run.font.bold = True

    for iface in interfaces:
        row = table.add_row().cells
        row[0].text = _extract_value(iface, "interface") or ""
        row[1].text = _extract_value(iface, "component_a") or ""
        row[2].text = _extract_value(iface, "component_b") or ""
        row[3].text = _extract_value(iface, "protocol") or ""
        row[4].text = _extract_value(iface, "authentication") or ""
        row[5].text = _extract_value(iface, "data_exchanged") or ""


def _append_environment_block(document: Document, title: str, html_content: Optional[str]) -> None:
    if not html_content:
        return
    subheading = document.add_paragraph()
    subheading_run = subheading.add_run(title)
    subheading_run.font.size = Pt(13)
    subheading_run.font.bold = True
    subheading.space_before = Pt(10)
    subheading.space_after = Pt(4)
    append_html_to_document(document, html_content)


def _extract_value(payload: object, field: str) -> Optional[str]:
    if isinstance(payload, dict):
        value = payload.get(field)
    else:
        value = getattr(payload, field, None)
    if value and isinstance(value, str) and value.strip():
        return value
    return None


def _strip_html(value: Optional[str]) -> str:
    if not value:
        return ""
    return re.sub(r"<[^>]+>", " ", value).strip()


def _append_product_user_description_section(document: Document, payload: Optional[object]) -> None:
    """Append Section 5.2.5 Product User Description."""
    from docx.shared import RGBColor
    
    if not payload:
        return

    # Start 5.2.5 Product User Description on a new page
    document.add_page_break()

    heading = document.add_paragraph()
    heading_run = heading.add_run("5.2.5 Product User Description")
    heading_run.font.size = Pt(16)
    heading_run.font.bold = True
    heading.space_after = Pt(6)

    reference = document.add_paragraph("[Reference: Clause 6.2.1.6 - Product user description]")
    reference.runs[0].font.bold = True
    reference.space_after = Pt(10)

    # User Description
    user_desc_html = _extract_value(payload, "user_description_html")
    if user_desc_html:
        append_html_to_document(document, user_desc_html)

    # RDPS Considerations
    no_rdps = getattr(payload, "no_rdps", False)
    rdps_label = document.add_paragraph()
    rdps_run = rdps_label.add_run("RDPS Considerations:")
    rdps_run.font.size = Pt(13)
    rdps_run.font.bold = True
    rdps_label.space_before = Pt(10)
    rdps_label.space_after = Pt(4)

    if no_rdps:
        no_rdps_para = document.add_paragraph("This product does not rely on any RDPS.")
        no_rdps_para.runs[0].font.italic = True
        no_rdps_para.space_after = Pt(8)
    else:
        rdps_html = _extract_value(payload, "rdps_considerations_html")
        if rdps_html:
            append_html_to_document(document, rdps_html)
        else:
            no_content_para = document.add_paragraph("No RDPS considerations specified.")
            no_content_para.runs[0].font.italic = True
            no_content_para.space_after = Pt(8)

    # Evidence Reference
    _append_evidence_tracker(document, getattr(payload, "evidence_entries", None))


def _append_product_context_assessment_section(document: Document, payload: Optional[object]) -> None:
    """Append Section 5.3 Product Context Assessment Summary."""
    from docx.shared import RGBColor
    
    if not payload:
        return

    # Start 5.3 Product Context Assessment on a new page
    document.add_page_break()

    heading = document.add_paragraph()
    heading_run = heading.add_run("5.3 Product Context Assessment Summary")
    heading_run.font.size = Pt(18)
    heading_run.font.bold = True
    heading.space_after = Pt(6)

    reference = document.add_paragraph("[Reference: Clause 6.2]")
    reference.runs[0].font.bold = True
    reference.space_after = Pt(10)

    # Overall Verdict
    verdict = _extract_value(payload, "overall_verdict") or "not_assessed"
    verdict_label = OVERALL_VERDICT_LABELS.get(verdict, verdict.upper())
    
    verdict_heading = document.add_paragraph()
    verdict_run = verdict_heading.add_run("Overall Verdict:")
    verdict_run.font.size = Pt(13)
    verdict_run.font.bold = True
    verdict_heading.space_before = Pt(6)
    verdict_heading.space_after = Pt(4)

    verdict_para = document.add_paragraph()
    verdict_value_run = verdict_para.add_run(verdict_label)
    verdict_value_run.font.bold = True
    verdict_value_run.font.size = Pt(12)
    # Color code the verdict
    if verdict == "pass":
        verdict_value_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
    elif verdict == "fail":
        verdict_value_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
    elif verdict == "partial":
        verdict_value_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
    verdict_para.space_after = Pt(10)

    # Summary of Findings
    summary_html = _extract_value(payload, "summary_of_findings_html")
    if summary_html:
        summary_heading = document.add_paragraph()
        summary_run = summary_heading.add_run("Summary of Findings:")
        summary_run.font.size = Pt(13)
        summary_run.font.bold = True
        summary_heading.space_before = Pt(10)
        summary_heading.space_after = Pt(4)
        append_html_to_document(document, summary_html)

    # Assessment Checklist Table
    assessments = getattr(payload, "assessments", None) or []
    if assessments:
        _append_assessment_checklist_table(document, assessments)

    # Non-Conformities Table
    non_conformities = getattr(payload, "non_conformities", None) or []
    if non_conformities:
        _append_non_conformities_table(document, non_conformities)


def _append_assessment_checklist_table(document: Document, assessments: list) -> None:
    """Append the assessment checklist table."""
    checklist_heading = document.add_paragraph()
    checklist_run = checklist_heading.add_run("Assessment Checklist:")
    checklist_run.font.size = Pt(13)
    checklist_run.font.bold = True
    checklist_heading.space_before = Pt(10)
    checklist_heading.space_after = Pt(4)

    # Build a lookup map from assessments
    assessment_map = {}
    for a in assessments:
        req_id = _extract_value(a, "id")
        if req_id:
            assessment_map[req_id] = a

    # Create table with headers
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["ID", "Requirement", "Evidence", "Status", "Comments"]
    header_cells = table.rows[0].cells
    for idx, label in enumerate(headers):
        paragraph = header_cells[idx].paragraphs[0]
        run = paragraph.add_run(label)
        run.font.bold = True

    # Track current subsection for grouping
    current_subsection = None

    for req in PRODUCT_CONTEXT_REQUIREMENTS:
        req_id = req["id"]
        subsection = req["subsection"]
        description = req["description"]

        # Add subsection header row if subsection changed
        if subsection != current_subsection:
            current_subsection = subsection
            subsection_label = req["label"]
            header_row = table.add_row().cells
            # Merge cells for subsection header
            header_row[0].merge(header_row[4])
            header_para = header_row[0].paragraphs[0]
            header_run = header_para.add_run(f"{subsection} - {subsection_label}")
            header_run.font.bold = True
            header_run.font.italic = True

        # Get assessment data for this requirement
        assessment = assessment_map.get(req_id, {})
        evidence_ref = _extract_value(assessment, "evidence_ref_id") or ""
        status = _extract_value(assessment, "status") or "not_assessed"
        status_label = ASSESSMENT_STATUS_LABELS.get(status, status)
        comments_html = _extract_value(assessment, "comments_html") or ""
        comments_text = _strip_html(comments_html)
        # Truncate comments for table display
        if len(comments_text) > 100:
            comments_text = comments_text[:97] + "..."

        # Add requirement row
        row = table.add_row().cells
        row[0].text = req_id
        row[1].text = description
        row[2].text = evidence_ref
        row[3].text = status_label
        row[4].text = comments_text


def _append_non_conformities_table(document: Document, non_conformities: list) -> None:
    """Append the non-conformities table."""
    from docx.shared import RGBColor
    
    nc_heading = document.add_paragraph()
    nc_run = nc_heading.add_run("Non-Conformities:")
    nc_run.font.size = Pt(13)
    nc_run.font.bold = True
    nc_heading.space_before = Pt(10)
    nc_heading.space_after = Pt(4)

    # Create table with headers
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["NC ID", "Requirement", "Description", "Severity", "Corrective Action"]
    header_cells = table.rows[0].cells
    for idx, label in enumerate(headers):
        paragraph = header_cells[idx].paragraphs[0]
        run = paragraph.add_run(label)
        run.font.bold = True

    for nc in non_conformities:
        nc_id = _extract_value(nc, "id") or ""
        requirement_id = _extract_value(nc, "requirement_id") or ""
        description = _extract_value(nc, "description") or ""
        severity = _extract_value(nc, "severity") or "minor"
        severity_label = NC_SEVERITY_LABELS.get(severity, severity.title())
        corrective_action = _extract_value(nc, "corrective_action") or ""

        row = table.add_row().cells
        row[0].text = nc_id
        row[1].text = requirement_id
        row[2].text = description
        
        # Add severity with color coding
        severity_para = row[3].paragraphs[0]
        severity_run = severity_para.add_run(severity_label)
        if severity == "critical":
            severity_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            severity_run.font.bold = True
        elif severity == "major":
            severity_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
        
        row[4].text = corrective_action
