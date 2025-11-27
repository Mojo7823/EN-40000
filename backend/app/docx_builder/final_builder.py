"""Final combined CRA documentation builder."""
import uuid
from pathlib import Path
from docx import Document
from docx.shared import Pt

from .section_builders import create_base_document, add_documentation_intro_section, add_section_with_html
from .cover_builder import add_cover_to_document
from .html_converter import append_html_to_document
from .risk_management_builder import append_risk_management_section


def build_final_combined_document(payload, image_file: Path, output_dir: Path) -> Path:
    """
    Build the complete final CRA Documentation document from all sections.
    
    Args:
        payload: FinalPreviewRequest with all document data
        image_file: Optional path to cover image
        output_dir: Directory to save document
        
    Returns:
        Path to generated DOCX file
    """
    # Clear previous previews
    for existing in output_dir.glob("*.docx"):
        existing.unlink(missing_ok=True)
    
    document = create_base_document()
    
    # Page 1: Add cover page if provided
    if payload.cover_data:
        try:
            add_cover_to_document(document, payload.cover_data, image_file)
        except Exception:
            pass  # Skip if cover generation fails
    
    # Page 2: Documentation Introduction
    add_documentation_intro_section(document)
    
    # Add Documentation Introduction subsections
    if payload.st_reference_html:
        add_section_with_html(
            document,
            "1.1",
            "Documentation Reference",
            payload.st_reference_html
        )
    
    if payload.toe_reference_html:
        add_section_with_html(
            document,
            "1.2",
            "Product Reference",
            payload.toe_reference_html
        )
    
    if payload.toe_overview_html:
        add_section_with_html(
            document,
            "1.3",
            "Product Overview",
            payload.toe_overview_html
        )
    
    if payload.toe_description_html:
        add_section_with_html(
            document,
            "1.4",
            "Product Description",
            payload.toe_description_html
        )
    
    # Section 2: Security Problem Definition
    if payload.spd_html:
        add_section_with_html(
            document,
            "2.",
            "Security Problem Definition",
            payload.spd_html,
            heading_size=20,
            add_page_break=True
        )
    
    # Section 3: Conformance Claims
    if payload.conformance_claims_html:
        add_section_with_html(
            document,
            "3.",
            "Conformance Claims",
            payload.conformance_claims_html,
            heading_size=20,
            add_page_break=True
        )
    
    # Section 4: Security Objectives
    if payload.security_objectives_html:
        add_section_with_html(
            document,
            "4.",
            "Security Objectives",
            payload.security_objectives_html,
            heading_size=20,
            add_page_break=True
        )
    
    # Section 5: Risk Management Elements
    if payload.risk_management:
        product_name = "[Product Name]"
        if payload.cover_data:
            product_name = payload.cover_data.get("title") or payload.cover_data.get("deviceName") or product_name
        append_risk_management_section(document, payload.risk_management, product_name)
    
    # Section 6: Security Requirements (renumbered from 5)
    _add_security_requirements_section(document, payload)
    
    # Section 7: Product Summary Specification (TSS)
    if payload.tss_html:
        document.add_page_break()
        
        tss_heading = document.add_paragraph()
        tss_run = tss_heading.add_run("7. Product Summary Specification")
        tss_run.font.size = Pt(20)
        tss_run.font.bold = True
        tss_heading.space_before = Pt(12)
        tss_heading.space_after = Pt(8)
        
        intro_paragraph = document.add_paragraph(
            (
                "This section describes the Product security functions that satisfy the security functional requirements. "
                "The Product also includes additional relevant security functions which are also described in the following "
                "sections, as well as a mapping to the security functional requirements satisfied by the Product."
            )
        )
        intro_paragraph.space_after = Pt(12)
        
        append_html_to_document(document, payload.tss_html)
    
    # Save document
    filename = f"{uuid.uuid4().hex}.docx"
    output_path = output_dir / filename
    document.save(str(output_path))
    return output_path


def _add_security_requirements_section(document: Document, payload):
    """
    Add Security Requirements section (SFR and SAR).
    
    Args:
        document: Document to add to
        payload: FinalPreviewRequest with requirements data
    """
    security_section_added = False
    
    # Add SFR section
    if payload.sfr_preview_html or (payload.sfr_list and len(payload.sfr_list) > 0):
        document.add_page_break()
        
        security_heading = document.add_paragraph()
        security_run = security_heading.add_run("6. Security Requirements")
        security_run.font.size = Pt(20)
        security_run.font.bold = True
        security_heading.space_before = Pt(12)
        security_heading.space_after = Pt(8)
        security_section_added = True
        
        sfr_heading = document.add_paragraph()
        sfr_run = sfr_heading.add_run("6.1 Security Functional Requirements")
        sfr_run.font.size = Pt(18)
        sfr_run.font.bold = True
        sfr_heading.space_before = Pt(8)
        sfr_heading.space_after = Pt(12)
        
        if payload.sfr_preview_html:
            append_html_to_document(document, payload.sfr_preview_html)
        else:
            for sfr_item in payload.sfr_list:
                if sfr_item.get('preview'):
                    append_html_to_document(document, sfr_item['preview'])
                    document.add_paragraph().space_after = Pt(12)
    
    # Add SAR section
    if payload.sar_preview_html or (payload.sar_list and len(payload.sar_list) > 0):
        if not security_section_added:
            document.add_page_break()
            
            security_heading = document.add_paragraph()
            security_run = security_heading.add_run("6. Security Requirements")
            security_run.font.size = Pt(20)
            security_run.font.bold = True
            security_heading.space_before = Pt(12)
            security_heading.space_after = Pt(8)
            security_section_added = True
        
        sar_heading = document.add_paragraph()
        sar_run = sar_heading.add_run("6.2 Security Assurance Requirements")
        sar_run.font.size = Pt(18)
        sar_run.font.bold = True
        sar_heading.space_before = Pt(8)
        sar_heading.space_after = Pt(12)
        
        if payload.sar_preview_html:
            append_html_to_document(document, payload.sar_preview_html)
        else:
            # Add EAL if specified
            if payload.selected_eal:
                eal_para = document.add_paragraph()
                eal_run = eal_para.add_run(f"Evaluation Assurance Level: {payload.selected_eal}")
                eal_run.font.bold = True
                eal_para.space_after = Pt(12)
            
            for sar_item in payload.sar_list:
                if sar_item.get('preview'):
                    append_html_to_document(document, sar_item['preview'])
                    document.add_paragraph().space_after = Pt(12)
