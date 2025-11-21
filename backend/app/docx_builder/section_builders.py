"""Document section builders for various CRA documentation sections."""
import uuid
from pathlib import Path
from docx import Document
from docx.shared import Mm, Pt

from .html_converter import append_html_to_document


def create_base_document() -> Document:
    """
    Create a base document with standard page settings.
    
    Returns:
        Configured Document object
    """
    document = Document()
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)
    return document


def build_html_preview_document(html_content: str, user_id: str, output_dir: Path) -> Path:
    """
    Build a simple HTML preview document.
    
    Args:
        html_content: HTML content to convert
        user_id: User identifier
        output_dir: Directory to save document
        
    Returns:
        Path to generated DOCX file
    """
    # Clear previous previews
    for existing in output_dir.glob("*.docx"):
        existing.unlink(missing_ok=True)
    
    document = create_base_document()
    append_html_to_document(document, html_content)
    
    filename = f"{uuid.uuid4().hex}.docx"
    output_path = output_dir / filename
    document.save(str(output_path))
    return output_path


def build_tss_preview_document(html_content: str, user_id: str, output_dir: Path) -> Path:
    """
    Build Product Summary Specification preview document.
    
    Note: TSS = TOE Summary Specification (Common Criteria legacy term)
    TOE = Target of Evaluation (now: Product)
    
    Args:
        html_content: HTML content for TSS
        user_id: User identifier
        output_dir: Directory to save document
        
    Returns:
        Path to generated DOCX file
    """
    # Clear previous previews
    for existing in output_dir.glob("*.docx"):
        existing.unlink(missing_ok=True)
    
    document = create_base_document()
    
    # Add section heading
    heading = document.add_paragraph()
    heading_run = heading.add_run("6. Product Summary Specification")
    heading_run.font.size = Pt(20)
    heading_run.font.bold = True
    heading.space_after = Pt(8)
    
    # Add introduction
    intro_paragraph = document.add_paragraph(
        (
            "This section describes the Product security functions that satisfy the technical requirements. "
            "The Product also includes additional relevant security functions which are also described in the following "
            "sections, as well as a mapping to the technical requirements satisfied by the Product."
        )
    )
    intro_paragraph.space_after = Pt(12)
    
    # Add HTML content
    append_html_to_document(document, html_content)
    
    filename = f"{uuid.uuid4().hex}.docx"
    output_path = output_dir / filename
    document.save(str(output_path))
    return output_path


def add_documentation_intro_section(document: Document, intro_text: str = None):
    """
    Add CRA Documentation Introduction section header.
    
    Args:
        document: Document to add to
        intro_text: Optional custom intro text
    """
    heading = document.add_paragraph()
    heading_run = heading.add_run("1. CRA Documentation Introduction")
    heading_run.font.size = Pt(20)
    heading_run.font.bold = True
    heading.space_after = Pt(12)
    
    # Default intro text
    if intro_text is None:
        intro_text = (
            "This section presents the following information required for CRA (Cyber Resilience Act) compliance:\n"
            "• Identifies the CRA Documentation and the Product\n"
            "• Specifies the documentation conventions\n"
            "• Describes the organization of the documentation"
        )
    
    intro_para = document.add_paragraph(intro_text)
    intro_para.space_after = Pt(12)


def add_section_with_html(
    document: Document,
    section_number: str,
    section_title: str,
    html_content: str,
    heading_size: int = 18,
    add_page_break: bool = False
):
    """
    Add a numbered section with HTML content.
    
    Args:
        document: Document to add to
        section_number: Section number (e.g., "1.1", "2")
        section_title: Section title
        html_content: HTML content for section
        heading_size: Font size for heading in points
        add_page_break: Whether to add page break before section
    """
    if add_page_break:
        document.add_page_break()
    
    heading = document.add_paragraph()
    heading_run = heading.add_run(f"{section_number} {section_title}")
    heading_run.font.size = Pt(heading_size)
    heading_run.font.bold = True
    heading.space_before = Pt(12)
    heading.space_after = Pt(8)
    
    append_html_to_document(document, html_content)
