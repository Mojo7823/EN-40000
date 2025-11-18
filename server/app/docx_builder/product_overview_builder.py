"""Product Overview section rendering."""
from typing import Any, List

from docx import Document
from docx.shared import Pt
from lxml import html as lxml_html, etree

from .html_converter import append_html_to_document

DESCRIPTION_PLACEHOLDER = (
    "[Provide a detailed description of the product highlighting physical, software, connectivity, user interface, and data processing characteristics.]"
)
DESCRIPTION_PAGE_CHAR_LIMIT = 2800
ARCHITECTURE_PAGE_CHAR_LIMIT = 2800
MIN_CHARS_BEFORE_BREAK = 1200


def append_product_overview_section(document: Document, overview_data: Any):
    """Append Section 2 - Product Overview."""
    document.add_page_break()

    heading = document.add_paragraph()
    heading_run = heading.add_run("2. Product Overview")
    heading_run.font.size = Pt(20)
    heading_run.font.bold = True
    heading.space_after = Pt(10)

    sub_heading = document.add_paragraph()
    sub_run = sub_heading.add_run("2.1 Product Description")
    sub_run.font.size = Pt(18)
    sub_run.font.bold = True
    sub_heading.space_after = Pt(6)

    reference = document.add_paragraph("[Reference: Clause 6.2 - Product Context]")
    reference.runs[0].font.bold = True
    reference.space_after = Pt(10)

    intro = document.add_paragraph(
        "Provide a detailed description of the product, including the following areas as applicable:"
    )
    intro.runs[0].font.italic = True
    intro.space_after = Pt(10)

    _add_guidance_line(document, "Physical Characteristics", "Describe hardware components, form factor, interfaces.")
    _add_guidance_line(document, "Software Characteristics", "Describe software architecture, programming languages, frameworks.")
    _add_guidance_line(document, "Network Connectivity", "Describe network capabilities, protocols, communication mechanisms.")
    _add_guidance_line(document, "User Interface", "Describe how users interact with the product (GUI, API, CLI, physical controls).")
    _add_guidance_line(document, "Data Processing", "Describe what types of data the product processes, stores, or transmits.")

    description_html = _get_overview_value(overview_data, "product_description_html")
    if description_html:
        _append_html_with_soft_breaks(document, description_html, DESCRIPTION_PAGE_CHAR_LIMIT)
    else:
        placeholder = document.add_paragraph(DESCRIPTION_PLACEHOLDER)
        placeholder.runs[0].font.italic = True

    # Begin architecture overview on a fresh page when content exists
    document.add_page_break()
    _add_architecture_section(document, overview_data)
    _add_third_party_section(document, overview_data)


def _add_guidance_line(document: Document, title: str, description: str):
    paragraph = document.add_paragraph()
    title_run = paragraph.add_run(f"{title}: ")
    title_run.font.bold = True
    detail_run = paragraph.add_run(description)
    detail_run.font.italic = True
    paragraph.space_after = Pt(4)


def _get_overview_value(source: Any, key: str) -> str:
    if not source:
        return ""
    if isinstance(source, dict):
        value = source.get(key)
    else:
        value = getattr(source, key, None)
    if value is None:
        return ""
    return str(value).strip()


def _add_architecture_section(document: Document, overview_data: Any):
    sub_heading = document.add_paragraph()
    sub_run = sub_heading.add_run("2.2 Product Architecture Overview")
    sub_run.font.size = Pt(18)
    sub_run.font.bold = True
    sub_heading.space_after = Pt(6)

    reference = document.add_paragraph("[Reference: Clause 6.2.1.5 - Product architecture overview]")
    reference.runs[0].font.bold = True
    reference.space_after = Pt(6)

    intro = document.add_paragraph(
        "[Provide a high-level architectural description of the product, including key components, interactions, interfaces, and remote services.]"
    )
    intro.runs[0].font.italic = True
    intro.space_after = Pt(8)

    architecture_html = _get_overview_value(overview_data, "product_architecture_html")
    if architecture_html:
        _append_html_with_soft_breaks(document, architecture_html, ARCHITECTURE_PAGE_CHAR_LIMIT)
        return

    _add_architecture_placeholder(document)


def _add_architecture_placeholder(document: Document):
    _add_key_components_placeholder(document)
    _add_architecture_hint(
        document,
        "Component Interactions",
        "[Describe how components interact with each other]",
    )
    _add_architecture_hint(
        document,
        "External Interfaces",
        "[List and describe all external interfaces - network ports, APIs, physical connections]",
    )
    _add_architecture_hint(
        document,
        "Remote Data Processing Solutions (RDPS)",
        "[Describe any cloud services, remote servers, or backend systems that are part of the product]",
    )
    diagram = document.add_paragraph("[Include an architecture diagram if available - reference it here]")
    diagram.runs[0].font.italic = True
    diagram.space_after = Pt(6)
    _add_architecture_hint(
        document,
        "Evidence Reference",
        "[Document ID/Location where architecture documentation can be found]",
    )


def _add_key_components_placeholder(document: Document):
    heading = document.add_paragraph("Key Components:")
    heading.runs[0].font.bold = True
    heading.space_after = Pt(2)

    placeholders = [
        "[Component 1 name and purpose]",
        "[Component 2 name and purpose]",
        "[Component 3 name and purpose]",
    ]

    for index, text in enumerate(placeholders, start=1):
        line = document.add_paragraph(f"{index}. {text}")
        line.paragraph_format.left_indent = Pt(12)
        run = line.runs[0]
        run.font.italic = True
        line.space_after = Pt(2)


def _add_architecture_hint(document: Document, title: str, hint: str):
    paragraph = document.add_paragraph()
    title_run = paragraph.add_run(f"{title}:")
    title_run.font.bold = True
    paragraph.space_after = Pt(2)

    detail = document.add_paragraph(hint)
    detail.paragraph_format.left_indent = Pt(12)
    detail.runs[0].font.italic = True
    detail.space_after = Pt(6)


def _add_third_party_section(document: Document, overview_data: Any):
    section_data = _get_nested_value(overview_data, "third_party_components") or {}
    document.add_page_break()
    heading = document.add_paragraph()
    run = heading.add_run("2.3 Third-Party Components")
    run.font.size = Pt(18)
    run.font.bold = True
    heading.space_after = Pt(4)

    reference = document.add_paragraph("[Reference: Clause 7.11 - Third-party component cybersecurity management]")
    reference.runs[0].font.bold = True
    reference.space_after = Pt(6)

    intro = document.add_paragraph(
        "The component of third-party software that has been used on this product are listed in this table below:"
    )
    intro.runs[0].font.italic = True
    intro.space_after = Pt(6)

    entries = _extract_third_party_entries(section_data)
    if entries:
        _render_components_table(document, entries)
    else:
        placeholder = document.add_paragraph(
            "[List all third-party components (hardware and software) integrated into the product]"
        )
        placeholder.runs[0].font.italic = True

    management_heading = document.add_paragraph("Third-Party Component Management Approach:")
    management_heading.runs[0].font.bold = True
    management_text = _get_overview_value(section_data, "management_approach_html")
    if management_text:
        append_html_to_document(document, management_text)
    else:
        management_placeholder = document.add_paragraph(
            "[Describe how third-party components are selected, evaluated, monitored, and updated]"
        )
        management_placeholder.runs[0].font.italic = True

    evidence_heading = document.add_paragraph("Evidence Reference:")
    evidence_heading.runs[0].font.bold = True
    evidence_text = _get_overview_value(section_data, "evidence_reference_html")
    if evidence_text:
        append_html_to_document(document, evidence_text)
    else:
        evidence_placeholder = document.add_paragraph(
            "[Document ID/Location where SBOM and component management records can be found]"
        )
        evidence_placeholder.runs[0].font.italic = True


def _get_nested_value(source: Any, key: str):
    if not source:
        return None
    if isinstance(source, dict):
        return source.get(key)
    return getattr(source, key, None)


def _extract_third_party_entries(section_data: Any) -> List[dict]:
    raw_entries = _get_nested_value(section_data, "entries")
    if not raw_entries:
        return []
    extracted = []
    for item in raw_entries:
        if isinstance(item, dict):
            extracted.append(item)
        else:
            extracted.append(
                {
                    "component_name": getattr(item, "component_name", None),
                    "component_type": getattr(item, "component_type", None),
                    "version": getattr(item, "version", None),
                    "supplier": getattr(item, "supplier", None),
                    "purpose": getattr(item, "purpose", None),
                    "license": getattr(item, "license", None),
                }
            )
    return extracted


def _render_components_table(document: Document, entries: List[dict]):
    headers = ["Component Name", "Type", "Version", "Supplier", "Purpose", "License"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, label in enumerate(headers):
        paragraph = header_cells[idx].paragraphs[0]
        run = paragraph.add_run(label)
        run.font.bold = True

    for entry in entries:
        row_cells = table.add_row().cells
        row_cells[0].text = entry.get("component_name") or ""
        row_cells[1].text = entry.get("component_type") or ""
        row_cells[2].text = entry.get("version") or ""
        row_cells[3].text = entry.get("supplier") or ""
        row_cells[4].text = entry.get("purpose") or ""
        row_cells[5].text = entry.get("license") or ""


def _append_html_with_soft_breaks(document: Document, html_content: str, char_limit: int):
    chunks = _chunk_html_content(html_content, char_limit)
    if not chunks:
        return

    for index, chunk in enumerate(chunks):
        append_html_to_document(document, chunk)
        if index < len(chunks) - 1:
            document.add_page_break()


def _chunk_html_content(html_content: str, char_limit: int) -> List[str]:
    if not html_content:
        return []

    try:
        fragment = lxml_html.fragment_fromstring(html_content, create_parent=True)
    except Exception:
        return [html_content]

    chunks: List[str] = []
    current: List[str] = []
    running_length = 0

    for child in fragment:
        serialized = etree.tostring(child, encoding="unicode", with_tail=False)
        text_content = (child.text_content() or "").strip()
        text_length = len(text_content)

        if current and running_length >= char_limit and text_length > 0:
            chunks.append("".join(current))
            current = [serialized]
            running_length = text_length
            continue

        if running_length + text_length > char_limit and running_length >= MIN_CHARS_BEFORE_BREAK:
            chunks.append("".join(current))
            current = [serialized]
            running_length = text_length
        else:
            current.append(serialized)
            running_length += text_length

    if current:
        chunks.append("".join(current))

    return chunks or [html_content]
