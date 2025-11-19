"""Cover page document builder."""
import uuid
from pathlib import Path
from typing import Any, Iterable, Optional

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Mm, Pt

from ..utils.formatters import format_cover_date
from .introduction_sections import IntroductionSectionsRenderer
from .product_overview_builder import append_product_overview_section
from .conformance_claim_builder import append_conformance_claim_section
from .document_convention_builder import render_document_convention_to_document

COVER_HEADER_TEXT = "EN 40000-1-2-2025 Conformity Assessment"


def build_cover_document(
    payload,
    image_file: Optional[Path] = None,
    output_dir: Optional[Path] = None,
) -> Path:
    """Generate a cover preview for the supplied payload."""
    for existing in output_dir.glob("*.docx"):
        existing.unlink(missing_ok=True)

    document = _create_a4_document()
    renderer = CoverDocumentRenderer(document)
    renderer.render_cover_page(payload, image_file)
    renderer.render_introduction_sections(
        introduction_payload=getattr(payload, "introduction", None),
        purpose_scope_payload=getattr(payload, "purpose_scope", None),
        product_identification_payload=getattr(payload, "product_identification", None),
        manufacturer_information_payload=getattr(payload, "manufacturer_information", None),
    )
    append_product_overview_section(document, getattr(payload, "product_overview", None))
    append_conformance_claim_section(document, getattr(payload, "conformance_claim", None))
    
    # Add Document Convention section if present
    convention_payload = getattr(payload, "document_convention", None)
    if convention_payload:
        convention_dict = convention_payload.model_dump() if hasattr(convention_payload, 'model_dump') else convention_payload
        render_document_convention_to_document(
            document,
            terminology_entries=convention_dict.get('terminology_entries'),
            evidence_notation_html=convention_dict.get('evidence_notation_html'),
            requirement_notation_html=convention_dict.get('requirement_notation_html'),
            assessment_verdicts_html=convention_dict.get('assessment_verdicts_html'),
            start_on_new_page=True,
        )

    filename = f"{uuid.uuid4().hex}.docx"
    output_path = output_dir / filename
    document.save(str(output_path))
    return output_path


def add_cover_to_document(document: Document, cover_data: dict, image_file: Optional[Path] = None):
    """Insert the cover page layout into an existing document."""
    CoverDocumentRenderer(document).render_cover_page(cover_data, image_file)
    document.add_page_break()


def _create_a4_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)
    return document


class CoverDocumentRenderer:
    """Handles cover rendering and delegates intro sections."""

    def __init__(self, document: Document):
        self.document = document

    def render_cover_page(self, data: Any, image_file: Optional[Path]):
        self._add_cover_header()
        self._add_cover_title_and_description(data)
        self._add_cover_image(image_file)
        self._add_version_block(data)
        self._add_footer_block(data)

    def render_introduction_sections(
        self,
        introduction_payload: Any,
        purpose_scope_payload: Any,
        product_identification_payload: Any,
        manufacturer_information_payload: Any,
    ):
        IntroductionSectionsRenderer(self.document).render(
            introduction_payload,
            purpose_scope_payload,
            product_identification_payload,
            manufacturer_information_payload,
        )

    # ------------------------------------------------------------------
    # Cover rendering helpers
    # ------------------------------------------------------------------
    def _add_cover_header(self):
        self._add_centered_paragraph(
            COVER_HEADER_TEXT,
            size=15,
            bold=True,
            space_after=Pt(28),
            space_before=Pt(18),
        )

    def _add_cover_title_and_description(self, data: Any):
        title_text = (self._get_value(data, "title") or "").strip() or "CRA Documentation Title"
        self._add_centered_paragraph(title_text, size=28, bold=True, space_after=Pt(16))

        for line in self._split_lines(self._get_value(data, "description")):
            self._add_centered_paragraph(line, size=14, space_after=Pt(6))

        self._add_centered_paragraph("", space_after=Pt(10))

    def _add_cover_image(self, image_file: Optional[Path]):
        if not image_file:
            return
        image_paragraph = self.document.add_paragraph()
        image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = image_paragraph.add_run()
        run.add_picture(str(image_file), width=Mm(120))
        image_paragraph.space_after = Pt(22)

    def _add_version_block(self, data: Any):
        version_value = (self._get_value(data, "version") or "").strip() or "—"
        self._add_centered_paragraph(f"Version: {version_value}", size=14, space_after=Pt(12))

        revision_value = self._get_value(data, "revision")
        formatted_revision = format_cover_date(revision_value)
        self._add_centered_paragraph(f"Revision  : {formatted_revision}", size=14, space_after=Pt(28))

    def _add_footer_block(self, data: Any):
        section = self.document.sections[0]
        section.different_first_page_header_footer = True
        footer = section.first_page_footer

        # Remove default empty paragraphs
        for paragraph in list(footer.paragraphs):
            p = paragraph._element
            p.getparent().remove(p)

        def _add_footer_line(text: str, *, bold: bool = False, size: int = 13):
            paragraph = footer.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph.space_after = Pt(4)
            run = paragraph.add_run(text)
            run.font.size = Pt(size)
            run.font.bold = bold

        _add_footer_line("Document Prepared By", bold=True, size=14)

        manufacturer_lines = self._split_lines(self._get_value(data, "manufacturer"))
        address_lines = self._split_lines(self._get_value(data, "laboratory_address"))

        if not manufacturer_lines and not address_lines:
            _add_footer_line("—")
            return

        for line in manufacturer_lines + address_lines:
            _add_footer_line(line)

    def _add_centered_paragraph(
        self,
        text: str,
        *,
        size: int = 12,
        bold: bool = False,
        space_after: Pt = Pt(6),
        space_before: Optional[Pt] = None,
    ):
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if space_before is not None:
            paragraph.paragraph_format.space_before = space_before
        run = paragraph.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        paragraph.space_after = space_after
        return paragraph

    @staticmethod
    def _split_lines(value: Any) -> Iterable[str]:
        if not value:
            return []
        return [line.strip() for line in str(value).splitlines() if line.strip()]

    @staticmethod
    def _get_value(data: Any, key: str):
        if isinstance(data, dict):
            return data.get(key)
        return getattr(data, key, None)
