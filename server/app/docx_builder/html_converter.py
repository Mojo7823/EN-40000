"""HTML to DOCX conversion utilities."""
from io import BytesIO
from typing import Optional, Dict
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from lxml import html as lxml_html

from ..utils.style_parser import (
    collect_inline_styles, 
    merge_styles, 
    apply_styles_to_run,
    parse_margin_left,
    parse_text_alignment,
    parse_image_alignment,
)
from ..utils.converters import px_to_mm
from ..utils.dimension_parser import extract_dimension_px, extract_table_column_widths
from ..utils.image_handler import decode_base64_image


def paragraph_has_content(paragraph) -> bool:
    """Check if paragraph has any text content."""
    if (paragraph.text or "").strip():
        return True
    for run in paragraph.runs:
        if (getattr(run, "text", "") or "").strip():
            return True
    return False


def append_inline_content(
    paragraph,
    element,
    inherited_styles: Optional[Dict] = None,
    suppress_leading_break: bool = False,
):
    """
    Append inline HTML content to a docx paragraph.
    
    Handles: text, images, line breaks, formatting tags (bold, italic, etc.),
    lists (ul, ol), and nested elements.
    
    Args:
        paragraph: python-docx Paragraph object
        element: lxml HTML element
        inherited_styles: Style dictionary from parent elements
    """
    styles = collect_inline_styles(element)
    combined_styles = merge_styles(inherited_styles or {
        "bold": False,
        "italic": False,
        "underline": False,
        "strike": False,
        "color": None,
        "size": None,
    }, styles)
    
    tag = (element.tag or "").lower()
    
    # Handle images
    if tag == "img":
        image_data = decode_base64_image(element.get("src", ""))
        if image_data:
            run = paragraph.add_run()
            apply_styles_to_run(run, combined_styles)
            image_stream = BytesIO(image_data)
            
            # Get dimensions
            width_px = extract_dimension_px(element, "width")
            height_px = extract_dimension_px(element, "height")
            
            kwargs = {}
            if width_px:
                kwargs["width"] = Mm(px_to_mm(width_px))
            if not kwargs and height_px:
                kwargs["height"] = Mm(px_to_mm(height_px))
            
            try:
                run.add_picture(image_stream, **kwargs)
            except Exception:
                pass  # Skip invalid images
        
        # Handle tail text
        tail = element.tail or ""
        if tail:
            run = paragraph.add_run(tail.replace("\xa0", " "))
            apply_styles_to_run(run, combined_styles)
        return
    
    # Handle line breaks
    if tag == "br":
        run = paragraph.add_run()
        run.add_break()
        tail = element.tail or ""
        if tail:
            tail_run = paragraph.add_run(tail.replace("\xa0", " "))
            apply_styles_to_run(tail_run, combined_styles)
        return
    
    # Handle block-level elements inline (p, div)
    if tag in {"p", "div"}:
        if paragraph_has_content(paragraph) and not suppress_leading_break:
            paragraph.add_run().add_break()
        
        text_content = element.text or ""
        if text_content:
            run = paragraph.add_run(text_content.replace("\xa0", " "))
            apply_styles_to_run(run, combined_styles)
        
        for child in element:
            append_inline_content(paragraph, child, combined_styles)
            tail = child.tail or ""
            if tail:
                tail_run = paragraph.add_run(tail.replace("\xa0", " "))
                apply_styles_to_run(tail_run, combined_styles)
        
        element_tail = element.tail or ""
        if element_tail:
            tail_run = paragraph.add_run(element_tail.replace("\xa0", " "))
            apply_styles_to_run(tail_run, combined_styles)
        return
    
    # Handle lists (ul, ol)
    if tag in {"ul", "ol"}:
        items = [child for child in element if (child.tag or "").lower() == "li"]
        for idx, child in enumerate(items, start=1):
            if paragraph_has_content(paragraph):
                paragraph.add_run().add_break()
            
            prefix = "• " if tag == "ul" else f"{idx}. "
            prefix_run = paragraph.add_run(prefix)
            apply_styles_to_run(prefix_run, combined_styles)
            append_inline_content(paragraph, child, combined_styles)
        
        list_tail = element.tail or ""
        if list_tail:
            tail_run = paragraph.add_run(list_tail.replace("\xa0", " "))
            apply_styles_to_run(tail_run, combined_styles)
        return
    
    # Handle list items
    if tag == "li":
        text_content = element.text or ""
        if text_content:
            run = paragraph.add_run(text_content.replace("\xa0", " "))
            apply_styles_to_run(run, combined_styles)
        
        for child in element:
            append_inline_content(paragraph, child, combined_styles, suppress_leading_break=True)
            tail = child.tail or ""
            if tail:
                tail_run = paragraph.add_run(tail.replace("\xa0", " "))
                apply_styles_to_run(tail_run, combined_styles)
        
        element_tail = element.tail or ""
        if element_tail:
            tail_run = paragraph.add_run(element_tail.replace("\xa0", " "))
            apply_styles_to_run(tail_run, combined_styles)
        return
    
    # Default: handle as inline text with formatting
    text = element.text or ""
    if text:
        run = paragraph.add_run(text.replace("\xa0", " "))
        apply_styles_to_run(run, combined_styles)
    
    for child in element:
        append_inline_content(paragraph, child, combined_styles)


def append_block_element(document: Document, element, inherited_indent: Optional[float] = None):
    """
    Append block-level HTML element to docx document.
    
    Handles: headings (h1-h6), paragraphs, divs, lists, tables, line breaks.
    
    Args:
        document: python-docx Document object
        element: lxml HTML element
        inherited_indent: Inherited left indent in points
    """
    tag = (element.tag or "").lower()
    style_attr = element.get("style")
    margin_left = parse_margin_left(style_attr)
    indent = margin_left if margin_left is not None else inherited_indent
    
    # Heading styles
    heading_map = {
        "h1": Pt(24),
        "h2": Pt(20),
        "h3": Pt(18),
        "h4": Pt(16),
        "h5": Pt(14),
        "h6": Pt(12),
    }
    
    if tag in heading_map:
        paragraph = document.add_paragraph()
        if indent:
            paragraph.paragraph_format.left_indent = Pt(indent)
        
        # Apply text alignment
        alignment = parse_text_alignment(element)
        if alignment is not None:
            paragraph.alignment = alignment
        
        base_styles = {
            "bold": True,
            "italic": False,
            "underline": False,
            "strike": False,
            "color": None,
            "size": heading_map[tag],
        }
        append_inline_content(paragraph, element, base_styles)
        return
    
    # Paragraph
    if tag == "p":
        # Check if paragraph is empty
        text_content = (element.text or "") + "".join(
            (child.text or "") + (child.tail or "") for child in element
        )
        if not text_content.strip() and not element.findall("*"):
            document.add_paragraph()
            return
        
        paragraph = document.add_paragraph()
        if indent:
            paragraph.paragraph_format.left_indent = Pt(indent)
        
        # Apply text alignment from paragraph
        alignment = parse_text_alignment(element)
        if alignment is not None:
            paragraph.alignment = alignment
        else:
            # Check if paragraph contains a single image and use its alignment
            children = list(element)
            if len(children) == 1 and (children[0].tag or "").lower() == "img" and not (element.text or "").strip():
                img_alignment = parse_image_alignment(children[0])
                if img_alignment is not None:
                    paragraph.alignment = img_alignment
        
        append_inline_content(paragraph, element)
        return
    
    # Div / Section
    if tag in {"div", "section"}:
        child_indent = indent if indent is not None else inherited_indent
        
        text = (element.text or "").strip()
        if text:
            paragraph = document.add_paragraph(text)
            if child_indent:
                paragraph.paragraph_format.left_indent = Pt(child_indent)
        
        for child in element:
            append_block_element(document, child, child_indent)
        
        tail = (element.tail or "").strip()
        if tail:
            paragraph = document.add_paragraph(tail)
            if child_indent:
                paragraph.paragraph_format.left_indent = Pt(child_indent)
        return
    
    # Lists (ul, ol)
    if tag in {"ul", "ol"}:
        items = [child for child in element if (child.tag or "").lower() == "li"]
        for idx, child in enumerate(items, start=1):
            paragraph = document.add_paragraph()
            if indent:
                paragraph.paragraph_format.left_indent = Pt(indent)
            
            prefix = "• " if tag == "ul" else f"{idx}. "
            run = paragraph.add_run(prefix)
            append_inline_content(paragraph, child, {
                "bold": False,
                "italic": False,
                "underline": False,
                "strike": False,
                "color": None,
                "size": None,
            })
        return
    
    # Tables
    if tag == "table":
        rows = [row for row in element.findall(".//tr")]
        if not rows:
            return
        
        # Build table structure
        max_cells = 0
        table_rows = []
        for row in rows:
            cells = [cell for cell in row if (cell.tag or "").lower() in {"th", "td"}]
            table_rows.append(cells)
            max_cells = max(max_cells, len(cells))
        
        if max_cells == 0:
            return
        
        # Create docx table
        table = document.add_table(rows=len(table_rows), cols=max_cells)
        table.style = "Table Grid"
        
        # Fill table cells
        for row_index, cells in enumerate(table_rows):
            for col_index, cell in enumerate(cells):
                if col_index >= max_cells:
                    continue
                
                paragraph = table.cell(row_index, col_index).paragraphs[0]
                paragraph.text = ""
                append_inline_content(paragraph, cell)
                
                # Bold header cells
                if (cell.tag or "").lower() == "th":
                    for run in paragraph.runs:
                        run.bold = True
        
        # Set column widths
        column_widths = extract_table_column_widths(element, max_cells)
        for idx, width_px in enumerate(column_widths):
            if width_px is None or idx >= len(table.columns):
                continue
            
            width_mm = px_to_mm(width_px)
            column_length = Mm(width_mm)
            table.columns[idx].width = column_length
            for row in table.rows:
                row.cells[idx].width = column_length
        return
    
    # Image (block-level)
    if tag == "img":
        image_data = decode_base64_image(element.get("src", ""))
        if image_data:
            paragraph = document.add_paragraph()
            if indent:
                paragraph.paragraph_format.left_indent = Pt(indent)
            
            # Apply alignment for block-level images, including custom container styles
            alignment = parse_image_alignment(element)
            if alignment is not None:
                paragraph.alignment = alignment
            
            run = paragraph.add_run()
            image_stream = BytesIO(image_data)
            
            # Get dimensions
            width_px = extract_dimension_px(element, "width")
            height_px = extract_dimension_px(element, "height")
            
            kwargs = {}
            if width_px:
                kwargs["width"] = Mm(px_to_mm(width_px))
            if not kwargs and height_px:
                kwargs["height"] = Mm(px_to_mm(height_px))
            
            try:
                run.add_picture(image_stream, **kwargs)
            except Exception:
                pass  # Skip invalid images
        return
    
    # Line break
    if tag == "br":
        document.add_paragraph()
        return
    
    # Fallback: treat unknown elements as paragraphs
    paragraph = document.add_paragraph()
    if indent:
        paragraph.paragraph_format.left_indent = Pt(indent)
    append_inline_content(paragraph, element)


def append_html_to_document(document: Document, html_content: str):
    """
    Parse HTML and append to docx document.
    
    Args:
        document: python-docx Document object
        html_content: HTML string to parse and convert
    """
    if not html_content or not html_content.strip():
        return
    
    try:
        fragment = lxml_html.fragment_fromstring(html_content, create_parent=True)
    except (ValueError, TypeError):
        # Fallback: add as plain text if parsing fails
        paragraph = document.add_paragraph(html_content)
        return
    
    for child in fragment:
        append_block_element(document, child)
