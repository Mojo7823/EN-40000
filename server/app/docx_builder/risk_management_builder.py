"""Risk Management Elements section builder."""
from typing import Optional

from docx import Document
from docx.shared import Pt

from .html_converter import append_html_to_document


def append_risk_management_section(document: Document, payload: Optional[object]) -> None:
    """Append Section 5 - Risk Management Elements to the document."""
    if not payload:
        return

    general_html = _extract_value(payload, "general_approach_html")
    if not general_html:
        return

    # Add page break before section
    document.add_page_break()

    # Main section heading: 5. RISK MANAGEMENT ELEMENTS
    heading = document.add_paragraph()
    heading_run = heading.add_run("5. RISK MANAGEMENT ELEMENTS")
    heading_run.font.size = Pt(20)
    heading_run.font.bold = True
    heading.space_after = Pt(8)

    # Reference right after main heading
    reference = document.add_paragraph("[Reference: Clause 6.1 - General]")
    reference.runs[0].font.bold = True
    reference.space_after = Pt(10)

    # Subsection heading: 5.1 General Approach to Risk Management
    subheading = document.add_paragraph()
    subheading_run = subheading.add_run("5.1 General Approach to Risk Management")
    subheading_run.font.size = Pt(18)
    subheading_run.font.bold = True
    subheading.space_before = Pt(6)
    subheading.space_after = Pt(6)

    # Reference for subsection
    clause_reference = document.add_paragraph("[Reference: Clause 5 - Risk Management Elements]")
    clause_reference.runs[0].font.bold = True
    clause_reference.space_after = Pt(10)

    # Append user-provided HTML content directly (no summary paragraph)
    append_html_to_document(document, general_html)


def _extract_value(payload: object, field: str) -> Optional[str]:
    if isinstance(payload, dict):
        value = payload.get(field)
    else:
        value = getattr(payload, field, None)
    if value and isinstance(value, str) and value.strip():
        return value
    return None
