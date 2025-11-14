import os
import base64
import re
import shutil
import tempfile
import time
import uuid
from datetime import datetime
from contextlib import asynccontextmanager
from io import BytesIO
from pathlib import Path
from typing import List, Optional

from fastapi import FastAPI, Depends, HTTPException, Query, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from sqlalchemy.orm import Session
from sqlalchemy import text

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Mm, Pt, RGBColor
from lxml import html as lxml_html

from .database import Base, engine, get_db
from .models import Component
from .schemas import ComponentCreate, ComponentOut, ComponentUpdate
from pydantic import BaseModel, Field, ConfigDict


@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup
    Base.metadata.create_all(bind=engine)
    yield
    # Shutdown (if needed)


app = FastAPI(title="CRA Tool API", lifespan=lifespan)

COVER_UPLOAD_ROOT = Path(os.getenv("COVER_UPLOAD_DIR", Path(tempfile.gettempdir()) / "cratool_cover_uploads"))
COVER_UPLOAD_ROOT.mkdir(parents=True, exist_ok=True)
COVER_DOCX_ROOT = Path(os.getenv("COVER_DOCX_DIR", Path(tempfile.gettempdir()) / "cratool_cover_docx"))
COVER_DOCX_ROOT.mkdir(parents=True, exist_ok=True)
SFR_DOCX_ROOT = Path(os.getenv("SFR_DOCX_DIR", Path(tempfile.gettempdir()) / "cratool_sfr_docx"))
SFR_DOCX_ROOT.mkdir(parents=True, exist_ok=True)
SAR_DOCX_ROOT = Path(os.getenv("SAR_DOCX_DIR", Path(tempfile.gettempdir()) / "cratool_sar_docx"))
SAR_DOCX_ROOT.mkdir(parents=True, exist_ok=True)
ST_INTRO_DOCX_ROOT = Path(os.getenv("ST_INTRO_DOCX_DIR", Path(tempfile.gettempdir()) / "cratool_stintro_docx"))
ST_INTRO_DOCX_ROOT.mkdir(parents=True, exist_ok=True)
SPD_DOCX_ROOT = Path(os.getenv("SPD_DOCX_DIR", Path(tempfile.gettempdir()) / "cratool_spd_docx"))
SPD_DOCX_ROOT.mkdir(parents=True, exist_ok=True)
SO_DOCX_ROOT = Path(os.getenv("SO_DOCX_DIR", Path(tempfile.gettempdir()) / "cratool_so_docx"))
SO_DOCX_ROOT.mkdir(parents=True, exist_ok=True)
TSS_DOCX_ROOT = Path(os.getenv("TSS_DOCX_DIR", Path(tempfile.gettempdir()) / "cratool_tss_docx"))
TSS_DOCX_ROOT.mkdir(parents=True, exist_ok=True)
FINAL_DOCX_ROOT = Path(os.getenv("FINAL_DOCX_DIR", Path(tempfile.gettempdir()) / "cratool_final_docx"))
FINAL_DOCX_ROOT.mkdir(parents=True, exist_ok=True)

USER_ID_PATTERN = re.compile(r"^[A-Za-z0-9_-]{1,64}$")


def get_user_upload_dir(user_id: str, *, create: bool = False) -> Path:
    if not USER_ID_PATTERN.match(user_id):
        raise HTTPException(status_code=400, detail="Invalid user identifier")

    user_dir = COVER_UPLOAD_ROOT / user_id
    if create:
        user_dir.mkdir(parents=True, exist_ok=True)
    return user_dir


def _get_preview_docx_dir(root: Path, user_id: str, *, create: bool = False) -> Path:
    if not USER_ID_PATTERN.match(user_id):
        raise HTTPException(status_code=400, detail="Invalid user identifier")

    user_dir = root / user_id
    if create:
        user_dir.mkdir(parents=True, exist_ok=True)
    return user_dir


def get_user_docx_dir(user_id: str, *, create: bool = False) -> Path:
    return _get_preview_docx_dir(COVER_DOCX_ROOT, user_id, create=create)


class CoverPreviewRequest(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    user_id: str = Field(..., alias="user_id")
    title: Optional[str] = None
    version: Optional[str] = None
    revision: Optional[str] = None
    description: Optional[str] = None
    manufacturer: Optional[str] = None
    date: Optional[str] = None
    image_path: Optional[str] = Field(None, alias="image_path")


class HtmlPreviewRequest(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    user_id: str = Field(..., alias="user_id")
    html_content: str = Field(..., alias="html_content")


class STIntroPreviewRequest(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    user_id: str = Field(..., alias="user_id")
    cover_data: Optional[dict] = None
    st_reference_html: Optional[str] = None
    toe_reference_html: Optional[str] = None
    toe_overview_html: Optional[str] = None
    toe_description_html: Optional[str] = None


class FinalPreviewRequest(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    user_id: str = Field(..., alias="user_id")
    cover_data: Optional[dict] = None
    st_reference_html: Optional[str] = None
    toe_reference_html: Optional[str] = None
    toe_overview_html: Optional[str] = None
    toe_description_html: Optional[str] = None
    conformance_claims_html: Optional[str] = None
    spd_html: Optional[str] = None
    security_objectives_html: Optional[str] = None
    tss_html: Optional[str] = None
    sfr_list: List[dict] = Field(default_factory=list)
    sar_list: List[dict] = Field(default_factory=list)
    selected_eal: Optional[str] = None
    sfr_preview_html: Optional[str] = None
    sar_preview_html: Optional[str] = None


def _resolve_uploaded_image_path(image_path: Optional[str], user_id: str) -> Optional[Path]:
    if not image_path:
        return None

    expected_prefix = f"/cover/uploads/{user_id}/"
    if not image_path.startswith(expected_prefix):
        raise HTTPException(status_code=400, detail="Invalid image reference for preview generation")

    filename = Path(image_path).name
    upload_dir = get_user_upload_dir(user_id, create=False)
    image_file = upload_dir / filename
    if not image_file.exists():
        raise HTTPException(status_code=400, detail="Referenced cover image could not be found")
    return image_file


def _format_cover_date(date_value: Optional[str]) -> str:
    if not date_value:
        return "—"

    try:
        parsed = datetime.fromisoformat(date_value)
        return parsed.strftime("%d %B %Y")
    except ValueError:
        return date_value


def _build_cover_document(payload: CoverPreviewRequest) -> Path:
    image_file = _resolve_uploaded_image_path(payload.image_path, payload.user_id)
    docx_dir = get_user_docx_dir(payload.user_id, create=True)

    # Clear previous previews for the user to avoid stale documents piling up
    for existing in docx_dir.glob("*.docx"):
        existing.unlink(missing_ok=True)

    document = Document()
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)

    if image_file:
        image_paragraph = document.add_paragraph()
        image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = image_paragraph.add_run()
        run.add_picture(str(image_file), width=Mm(120))
        image_paragraph.space_after = Pt(12)

    title_text = payload.title.strip() if payload.title else "CRA Documentation Title"
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_paragraph.add_run(title_text)
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_paragraph.space_after = Pt(12)

    if payload.description:
        description_paragraph = document.add_paragraph(payload.description.strip())
        description_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        description_paragraph.space_after = Pt(18)

    info_items = [
        ("Version", payload.version.strip() if payload.version else "—"),
        ("Revision", payload.revision.strip() if payload.revision else "—"),
        ("Manufacturer/Laboratory", payload.manufacturer.strip() if payload.manufacturer else "—"),
        ("Date", _format_cover_date(payload.date)),
    ]

    for label, value in info_items:
        paragraph = document.add_paragraph()
        paragraph.space_after = Pt(6)
        run_label = paragraph.add_run(f"{label}: ")
        run_label.font.bold = True
        paragraph.add_run(value)

    filename = f"{uuid.uuid4().hex}.docx"
    output_path = docx_dir / filename
    document.save(str(output_path))
    return output_path


def _px_to_points(px_value: float) -> float:
    # Approximate conversion assuming 96px = 72pt
    return px_value * 0.75


def _parse_margin_left(style: Optional[str]) -> Optional[float]:
    if not style:
        return None
    match = re.search(r"margin-left\s*:\s*([0-9.]+)px", style)
    if not match:
        return None
    try:
        return _px_to_points(float(match.group(1)))
    except ValueError:
        return None


def _parse_color(value: Optional[str]) -> Optional[RGBColor]:
    if not value:
        return None

    color = value.strip().lower()
    if color.startswith("#"):
        hex_value = color[1:]
        if len(hex_value) == 3:
            hex_value = "".join(ch * 2 for ch in hex_value)
        if len(hex_value) == 6:
            try:
                r = int(hex_value[0:2], 16)
                g = int(hex_value[2:4], 16)
                b = int(hex_value[4:6], 16)
                return RGBColor(r, g, b)
            except ValueError:
                return None
    elif color.startswith("rgb"):
        numbers = re.findall(r"[0-9]{1,3}", color)
        if len(numbers) >= 3:
            try:
                r, g, b = (min(255, max(0, int(num))) for num in numbers[:3])
                return RGBColor(r, g, b)
            except ValueError:
                return None
    return None


def _collect_inline_styles(element) -> dict:
    styles = {
        "bold": False,
        "italic": False,
        "underline": False,
        "strike": False,
        "color": None,
        "size": None,
    }

    tag = (element.tag or "").lower()
    if tag in {"strong", "b"}:
        styles["bold"] = True
    if tag in {"em", "i"}:
        styles["italic"] = True
    if tag in {"u", "ins"}:
        styles["underline"] = True
    if tag in {"s", "strike", "del"}:
        styles["strike"] = True

    style_attr = element.get("style", "")
    for rule in style_attr.split(";"):
        rule = rule.strip().lower()
        if not rule:
            continue
        if "bold" in rule:
            styles["bold"] = True
        if "italic" in rule:
            styles["italic"] = True
        if "underline" in rule:
            styles["underline"] = True
        if "line-through" in rule:
            styles["strike"] = True
        if rule.startswith("color"):
            parts = rule.split(":", 1)
            if len(parts) == 2:
                parsed = _parse_color(parts[1])
                if parsed:
                    styles["color"] = parsed

    color_attr = element.get("color")
    parsed_color = _parse_color(color_attr)
    if parsed_color:
        styles["color"] = parsed_color

    return styles


CSS_WIDTH_RE = re.compile(r"width\s*:\s*([0-9.]+)px", re.IGNORECASE)
CSS_HEIGHT_RE = re.compile(r"height\s*:\s*([0-9.]+)px", re.IGNORECASE)


def _apply_styles_to_run(run, styles: dict):
    if styles.get("bold"):
        run.bold = True
    if styles.get("italic"):
        run.italic = True
    if styles.get("underline"):
        run.underline = True
    if styles.get("strike"):
        run.strike = True
    color = styles.get("color")
    if color:
        run.font.color.rgb = color
    size = styles.get("size")
    if size:
        run.font.size = size


def _merge_styles(parent: dict, child: dict) -> dict:
    merged = parent.copy()
    for key, value in child.items():
        if key in {"color", "size"}:
            if value is not None:
                merged[key] = value
        elif value:
            merged[key] = True
    return merged


def _px_to_mm(px: float) -> float:
    return px * 0.264583


def _extract_dimension_px(element, attr_name: str) -> Optional[float]:
    if element is None:
        return None

    style = element.get("style", "")
    regex = CSS_WIDTH_RE if attr_name == "width" else CSS_HEIGHT_RE
    match = regex.search(style)
    if match:
        try:
            return float(match.group(1))
        except (TypeError, ValueError):
            pass

    attr_value = element.get(attr_name)
    if attr_value:
        try:
            return float(attr_value)
        except (TypeError, ValueError):
            pass

    if attr_name == "width":
        colwidth = element.get("data-colwidth")
        if colwidth:
            try:
                parts = [float(part) for part in colwidth.split(",") if part.strip()]
                if parts:
                    return parts[0]
            except (TypeError, ValueError):
                pass

    return None


def _decode_base64_image(src: str) -> Optional[bytes]:
    if not src:
        return None

    if src.startswith("data:image"):
        try:
            header, data = src.split(",", 1)
        except ValueError:
            return None
        if ";base64" not in header:
            return None
        try:
            return base64.b64decode(data)
        except Exception:
            return None

    return None


def _append_inline_content(paragraph, element, inherited_styles: Optional[dict] = None):
    styles = _collect_inline_styles(element)
    combined_styles = _merge_styles(inherited_styles or {
        "bold": False,
        "italic": False,
        "underline": False,
        "strike": False,
        "color": None,
        "size": None,
    }, styles)

    tag = (element.tag or "").lower()

    def paragraph_has_content(p) -> bool:
        if (p.text or "").strip():
            return True
        for run in p.runs:
            if (getattr(run, "text", "") or "").strip():
                return True
        return False

    if tag == "img":
        image_data = _decode_base64_image(element.get("src", ""))
        if image_data:
            run = paragraph.add_run()
            _apply_styles_to_run(run, combined_styles)
            image_stream = BytesIO(image_data)
            width_px = _extract_dimension_px(element, "width")
            height_px = _extract_dimension_px(element, "height")
            kwargs = {}
            if width_px:
                kwargs["width"] = Mm(_px_to_mm(width_px))
            if not kwargs and height_px:
                kwargs["height"] = Mm(_px_to_mm(height_px))
            try:
                run.add_picture(image_stream, **kwargs)
            except Exception:
                pass
        tail = element.tail or ""
        if tail:
            run = paragraph.add_run(tail.replace("\xa0", " "))
            _apply_styles_to_run(run, combined_styles)
        return

    if tag == "br":
        run = paragraph.add_run()
        run.add_break()
        tail = element.tail or ""
        if tail:
            tail_run = paragraph.add_run(tail.replace("\xa0", " "))
            _apply_styles_to_run(tail_run, combined_styles)
        return

    if tag in {"p", "div"}:
        if paragraph_has_content(paragraph):
            paragraph.add_run().add_break()
        text_content = element.text or ""
        if text_content:
            run = paragraph.add_run(text_content.replace("\xa0", " "))
            _apply_styles_to_run(run, combined_styles)
        for child in element:
            _append_inline_content(paragraph, child, combined_styles)
            tail = child.tail or ""
            if tail:
                tail_run = paragraph.add_run(tail.replace("\xa0", " "))
                _apply_styles_to_run(tail_run, combined_styles)
        element_tail = element.tail or ""
        if element_tail:
            tail_run = paragraph.add_run(element_tail.replace("\xa0", " "))
            _apply_styles_to_run(tail_run, combined_styles)
        return

    if tag in {"ul", "ol"}:
        items = [child for child in element if (child.tag or "").lower() == "li"]
        for idx, child in enumerate(items, start=1):
            if paragraph_has_content(paragraph):
                paragraph.add_run().add_break()
            prefix = "• " if tag == "ul" else f"{idx}. "
            prefix_run = paragraph.add_run(prefix)
            _apply_styles_to_run(prefix_run, combined_styles)
            _append_inline_content(paragraph, child, combined_styles)
        list_tail = element.tail or ""
        if list_tail:
            tail_run = paragraph.add_run(list_tail.replace("\xa0", " "))
            _apply_styles_to_run(tail_run, combined_styles)
        return

    if tag == "li":
        text_content = element.text or ""
        if text_content:
            run = paragraph.add_run(text_content.replace("\xa0", " "))
            _apply_styles_to_run(run, combined_styles)
        for child in element:
            _append_inline_content(paragraph, child, combined_styles)
            tail = child.tail or ""
            if tail:
                tail_run = paragraph.add_run(tail.replace("\xa0", " "))
                _apply_styles_to_run(tail_run, combined_styles)
        element_tail = element.tail or ""
        if element_tail:
            tail_run = paragraph.add_run(element_tail.replace("\xa0", " "))
            _apply_styles_to_run(tail_run, combined_styles)
        return

    text = element.text or ""
    if text:
        run = paragraph.add_run(text.replace("\xa0", " "))
        _apply_styles_to_run(run, combined_styles)

    for child in element:
        _append_inline_content(paragraph, child, combined_styles)
        # Note: child handlers are responsible for processing their own tails
        # so we don't process tail here to avoid duplication


def _append_block_element(document: Document, element, inherited_indent: Optional[float] = None):
    tag = (element.tag or "").lower()
    style_attr = element.get("style")
    margin_left = _parse_margin_left(style_attr)
    indent = margin_left if margin_left is not None else inherited_indent

    heading_map = {
        "h1": Pt(24),
        "h2": Pt(20),
        "h3": Pt(18),
        "h4": Pt(16),
        "h5": Pt(14),
        "h6": Pt(12),
    }

    if tag in heading_map:
        paragraph = document.add_paragraph()
        if indent:
            paragraph.paragraph_format.left_indent = Pt(indent)
        base_styles = {
            "bold": True,
            "italic": False,
            "underline": False,
            "strike": False,
            "color": None,
            "size": heading_map[tag],
        }
        _append_inline_content(paragraph, element, base_styles)
        return

    if tag == "p":
        text_content = (element.text or "") + "".join(
            (child.text or "") + (child.tail or "") for child in element
        )
        if not text_content.strip() and not element.findall("*"):
            document.add_paragraph()
            return

        paragraph = document.add_paragraph()
        if indent:
            paragraph.paragraph_format.left_indent = Pt(indent)
        _append_inline_content(paragraph, element)
        return

    if tag in {"div", "section"}:
        child_indent = indent if indent is not None else inherited_indent
        text = (element.text or "").strip()
        if text:
            paragraph = document.add_paragraph(text)
            if child_indent:
                paragraph.paragraph_format.left_indent = Pt(child_indent)
        for child in element:
            _append_block_element(document, child, child_indent)
        tail = (element.tail or "").strip()
        if tail:
            paragraph = document.add_paragraph(tail)
            if child_indent:
                paragraph.paragraph_format.left_indent = Pt(child_indent)
        return

    if tag in {"ul", "ol"}:
        items = [child for child in element if (child.tag or "").lower() == "li"]
        for idx, child in enumerate(items, start=1):
            paragraph = document.add_paragraph()
            if indent:
                paragraph.paragraph_format.left_indent = Pt(indent)
            prefix = "• " if tag == "ul" else f"{idx}. "
            run = paragraph.add_run(prefix)
            _append_inline_content(paragraph, child, {
                "bold": False,
                "italic": False,
                "underline": False,
                "strike": False,
                "color": None,
                "size": None,
            })
        return

    if tag == "table":
        rows = [row for row in element.findall(".//tr")]
        if not rows:
            return

        max_cells = 0
        table_rows = []
        for row in rows:
            cells = [cell for cell in row if (cell.tag or "").lower() in {"th", "td"}]
            table_rows.append(cells)
            max_cells = max(max_cells, len(cells))

        if max_cells == 0:
            return

        table = document.add_table(rows=len(table_rows), cols=max_cells)
        table.style = "Table Grid"

        for row_index, cells in enumerate(table_rows):
            for col_index, cell in enumerate(cells):
                if col_index >= max_cells:
                    continue
                paragraph = table.cell(row_index, col_index).paragraphs[0]
                paragraph.text = ""
                _append_inline_content(paragraph, cell)
                if (cell.tag or "").lower() == "th":
                    for run in paragraph.runs:
                        run.bold = True

        column_widths = _extract_table_column_widths(element, max_cells)
        for idx, width_px in enumerate(column_widths):
            if width_px is None or idx >= len(table.columns):
                continue
            width_mm = _px_to_mm(width_px)
            column_length = Mm(width_mm)
            table.columns[idx].width = column_length
            for row in table.rows:
                row.cells[idx].width = column_length
        return

    if tag == "br":
        document.add_paragraph()
        return

    # Fallback: treat unknown block elements as paragraphs.
    paragraph = document.add_paragraph()
    if indent:
        paragraph.paragraph_format.left_indent = Pt(indent)
    _append_inline_content(paragraph, element)


def _extract_table_column_widths(table_element, max_cols: int) -> List[Optional[float]]:
    widths: List[Optional[float]] = [None] * max_cols

    colgroup = table_element.find("colgroup")
    if colgroup is not None:
        for idx, col in enumerate(colgroup.findall("col")):
            if idx >= max_cols:
                break
            width_px = _extract_dimension_px(col, "width")
            if width_px:
                widths[idx] = width_px

    for row in table_element.findall(".//tr"):
        cells = [cell for cell in row if (cell.tag or "").lower() in {"th", "td"}]
        for idx, cell in enumerate(cells):
            if idx >= max_cols:
                break
            if widths[idx] is None:
                width_px = _extract_dimension_px(cell, "width")
                if width_px:
                    widths[idx] = width_px

    return widths


def _append_html_to_document(document: Document, html_content: str):
    if not html_content or not html_content.strip():
        return

    try:
        fragment = lxml_html.fragment_fromstring(html_content, create_parent=True)
    except (ValueError, TypeError):
        paragraph = document.add_paragraph(html_content)
        return

    for child in fragment:
        _append_block_element(document, child)


def _build_html_preview_document(html_content: str, user_id: str, root: Path) -> Path:
    docx_dir = _get_preview_docx_dir(root, user_id, create=True)

    for existing in docx_dir.glob("*.docx"):
        existing.unlink(missing_ok=True)

    document = Document()
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)

    _append_html_to_document(document, html_content)

    filename = f"{uuid.uuid4().hex}.docx"
    output_path = docx_dir / filename
    document.save(str(output_path))
    return output_path


def _build_tss_preview_document(html_content: str, user_id: str) -> Path:
    docx_dir = _get_preview_docx_dir(TSS_DOCX_ROOT, user_id, create=True)

    for existing in docx_dir.glob("*.docx"):
        existing.unlink(missing_ok=True)

    document = Document()
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)

    heading = document.add_paragraph()
    heading_run = heading.add_run("6. Product Summary Specification")
    heading_run.font.size = Pt(20)
    heading_run.font.bold = True
    heading.space_after = Pt(8)

    intro_paragraph = document.add_paragraph(
        (
            "This section describes the Product security functions that satisfy the security functional requirements. "
            "The Product also includes additional relevant security functions which are also described in the following "
            "sections, as well as a mapping to the security functional requirements satisfied by the Product."
        )
    )
    intro_paragraph.space_after = Pt(12)

    _append_html_to_document(document, html_content)

    filename = f"{uuid.uuid4().hex}.docx"
    output_path = docx_dir / filename
    document.save(str(output_path))
    return output_path


def _build_st_intro_combined_document(payload: STIntroPreviewRequest) -> Path:
    """Build a combined ST Introduction document from all sections."""
    docx_dir = _get_preview_docx_dir(ST_INTRO_DOCX_ROOT, payload.user_id, create=True)

    # Clear previous previews
    for existing in docx_dir.glob("*.docx"):
        existing.unlink(missing_ok=True)

    document = Document()
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)

    # Add cover page if provided
    if payload.cover_data:
        cover_dict = payload.cover_data
        image_path = cover_dict.get("image_path")
        
        # Add cover image if present
        if image_path:
            try:
                image_file = _resolve_uploaded_image_path(image_path, payload.user_id)
                if image_file:
                    image_paragraph = document.add_paragraph()
                    image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = image_paragraph.add_run()
                    run.add_picture(str(image_file), width=Mm(120))
                    image_paragraph.space_after = Pt(12)
            except:
                pass  # Skip if image not found
        
        # Add cover title
        title_text = cover_dict.get("title", "").strip() or "CRA Documentation Title"
        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title_paragraph.add_run(title_text)
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_paragraph.space_after = Pt(12)
        
        # Add cover description
        if cover_dict.get("description"):
            description_paragraph = document.add_paragraph(cover_dict["description"].strip())
            description_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            description_paragraph.space_after = Pt(18)
        
        # Add cover metadata
        info_items = [
            ("Version", cover_dict.get("version", "").strip() or "—"),
            ("Revision", cover_dict.get("revision", "").strip() or "—"),
            ("Manufacturer/Laboratory", cover_dict.get("manufacturer", "").strip() or "—"),
            ("Date", _format_cover_date(cover_dict.get("date"))),
        ]
        
        for label, value in info_items:
            paragraph = document.add_paragraph()
            paragraph.space_after = Pt(6)
            run_label = paragraph.add_run(f"{label}: ")
            run_label.font.bold = True
            paragraph.add_run(value)
        
        # Add page break after cover
        document.add_page_break()

    # Add main heading
    heading = document.add_paragraph()
    heading_run = heading.add_run("1. CRA Documentation Introduction")
    heading_run.font.size = Pt(20)
    heading_run.font.bold = True
    heading.space_after = Pt(12)
    
    # Add introduction text
    intro_text = (
        "This section presents the following information required for a CRA (Cyber Resilience Act) (CC) evaluation:\n"
        "• Identifies the CRA Documentation (ST) and the Target of Evaluation (Product)\n"
        "• Specifies the security target conventions,\n"
        "• Describes the organization of the security target"
    )
    intro_para = document.add_paragraph(intro_text)
    intro_para.space_after = Pt(12)

    # Add ST Reference section
    if payload.st_reference_html:
        st_ref_heading = document.add_paragraph()
        st_ref_run = st_ref_heading.add_run("1.1 ST Reference")
        st_ref_run.font.size = Pt(18)
        st_ref_run.font.bold = True
        st_ref_heading.space_before = Pt(12)
        st_ref_heading.space_after = Pt(8)
        
        _append_html_to_document(document, payload.st_reference_html)

    # Add Product Reference section
    if payload.toe_reference_html:
        toe_ref_heading = document.add_paragraph()
        toe_ref_run = toe_ref_heading.add_run("1.2 Product Reference")
        toe_ref_run.font.size = Pt(18)
        toe_ref_run.font.bold = True
        toe_ref_heading.space_before = Pt(12)
        toe_ref_heading.space_after = Pt(8)
        
        _append_html_to_document(document, payload.toe_reference_html)

    # Add Product Overview section
    if payload.toe_overview_html:
        toe_overview_heading = document.add_paragraph()
        toe_overview_run = toe_overview_heading.add_run("1.3 Product Overview")
        toe_overview_run.font.size = Pt(18)
        toe_overview_run.font.bold = True
        toe_overview_heading.space_before = Pt(12)
        toe_overview_heading.space_after = Pt(8)
        
        _append_html_to_document(document, payload.toe_overview_html)

    # Add Product Description section
    if payload.toe_description_html:
        toe_desc_heading = document.add_paragraph()
        toe_desc_run = toe_desc_heading.add_run("1.4 Product Description")
        toe_desc_run.font.size = Pt(18)
        toe_desc_run.font.bold = True
        toe_desc_heading.space_before = Pt(12)
        toe_desc_heading.space_after = Pt(8)
        
        _append_html_to_document(document, payload.toe_description_html)

    filename = f"{uuid.uuid4().hex}.docx"
    output_path = docx_dir / filename
    document.save(str(output_path))
    return output_path


def _build_final_combined_document(payload: FinalPreviewRequest) -> Path:
    """Build the complete final CRA Documentation document from all sections."""
    docx_dir = _get_preview_docx_dir(FINAL_DOCX_ROOT, payload.user_id, create=True)

    # Clear previous previews
    for existing in docx_dir.glob("*.docx"):
        existing.unlink(missing_ok=True)

    document = Document()
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)

    # Page 1: Add cover page if provided
    if payload.cover_data:
        cover_dict = payload.cover_data
        image_path = cover_dict.get("image_path")
        
        # Add cover image if present
        if image_path:
            try:
                image_file = _resolve_uploaded_image_path(image_path, payload.user_id)
                if image_file:
                    image_paragraph = document.add_paragraph()
                    image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = image_paragraph.add_run()
                    run.add_picture(str(image_file), width=Mm(120))
                    image_paragraph.space_after = Pt(12)
            except:
                pass  # Skip if image not found
        
        # Add cover title
        title_text = cover_dict.get("title", "").strip() or "CRA Documentation Title"
        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title_paragraph.add_run(title_text)
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_paragraph.space_after = Pt(12)
        
        # Add cover description
        if cover_dict.get("description"):
            description_paragraph = document.add_paragraph(cover_dict["description"].strip())
            description_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            description_paragraph.space_after = Pt(18)
        
        # Add cover metadata
        info_items = [
            ("Version", cover_dict.get("version", "").strip() or "—"),
            ("Revision", cover_dict.get("revision", "").strip() or "—"),
            ("Manufacturer/Laboratory", cover_dict.get("manufacturer", "").strip() or "—"),
            ("Date", _format_cover_date(cover_dict.get("date"))),
        ]
        
        for label, value in info_items:
            paragraph = document.add_paragraph()
            paragraph.space_after = Pt(6)
            run_label = paragraph.add_run(f"{label}: ")
            run_label.font.bold = True
            paragraph.add_run(value)
        
        # Add page break after cover
        document.add_page_break()

    # Page 2: Add ST Introduction heading
    heading = document.add_paragraph()
    heading_run = heading.add_run("1. CRA Documentation Introduction")
    heading_run.font.size = Pt(20)
    heading_run.font.bold = True
    heading.space_after = Pt(12)

    intro_text = (
        "This section presents the following information required for a CRA (Cyber Resilience Act) (CC) evaluation:\n"
        "• Identifies the CRA Documentation (ST) and the Target of Evaluation (Product)\n"
        "• Specifies the security target conventions,\n"
        "• Describes the organization of the security target"
    )
    intro_para = document.add_paragraph(intro_text)
    intro_para.space_after = Pt(12)

    # Add ST Reference section
    if payload.st_reference_html:
        st_ref_heading = document.add_paragraph()
        st_ref_run = st_ref_heading.add_run("1.1 ST Reference")
        st_ref_run.font.size = Pt(18)
        st_ref_run.font.bold = True
        st_ref_heading.space_before = Pt(12)
        st_ref_heading.space_after = Pt(8)
        _append_html_to_document(document, payload.st_reference_html)

    # Page 3: Add Product Reference section
    if payload.toe_reference_html:
        toe_ref_heading = document.add_paragraph()
        toe_ref_run = toe_ref_heading.add_run("1.2 Product Reference")
        toe_ref_run.font.size = Pt(18)
        toe_ref_run.font.bold = True
        toe_ref_heading.space_before = Pt(12)
        toe_ref_heading.space_after = Pt(8)
        _append_html_to_document(document, payload.toe_reference_html)

    # Page 4: Add Product Overview section
    if payload.toe_overview_html:
        toe_overview_heading = document.add_paragraph()
        toe_overview_run = toe_overview_heading.add_run("1.3 Product Overview")
        toe_overview_run.font.size = Pt(18)
        toe_overview_run.font.bold = True
        toe_overview_heading.space_before = Pt(12)
        toe_overview_heading.space_after = Pt(8)
        _append_html_to_document(document, payload.toe_overview_html)

    # Page 5: Add Product Description section
    if payload.toe_description_html:
        toe_desc_heading = document.add_paragraph()
        toe_desc_run = toe_desc_heading.add_run("1.4 Product Description")
        toe_desc_run.font.size = Pt(18)
        toe_desc_run.font.bold = True
        toe_desc_heading.space_before = Pt(12)
        toe_desc_heading.space_after = Pt(8)
        _append_html_to_document(document, payload.toe_description_html)

    # Page 6: Add Security Problem Definition section
    if payload.spd_html:
        document.add_page_break()
        spd_heading = document.add_paragraph()
        spd_run = spd_heading.add_run("2. Security Problem Definition")
        spd_run.font.size = Pt(20)
        spd_run.font.bold = True
        spd_heading.space_before = Pt(12)
        spd_heading.space_after = Pt(12)
        _append_html_to_document(document, payload.spd_html)

    # Page 7: Add Conformance Claims section
    if payload.conformance_claims_html:
        document.add_page_break()
        conf_heading = document.add_paragraph()
        conf_run = conf_heading.add_run("3. Conformance Claims")
        conf_run.font.size = Pt(20)
        conf_run.font.bold = True
        conf_heading.space_before = Pt(12)
        conf_heading.space_after = Pt(8)
        _append_html_to_document(document, payload.conformance_claims_html)

    # Page 8: Add Security Objectives section
    if payload.security_objectives_html:
        document.add_page_break()
        so_heading = document.add_paragraph()
        so_run = so_heading.add_run("4. Security Objectives")
        so_run.font.size = Pt(20)
        so_run.font.bold = True
        so_heading.space_before = Pt(12)
        so_heading.space_after = Pt(8)
        _append_html_to_document(document, payload.security_objectives_html)

    # Page 8: Add Security Requirements sections
    security_section_added = False

    if payload.sfr_preview_html or (payload.sfr_list and len(payload.sfr_list) > 0):
        document.add_page_break()
        security_heading = document.add_paragraph()
        security_run = security_heading.add_run("5. Security Requirements")
        security_run.font.size = Pt(20)
        security_run.font.bold = True
        security_heading.space_before = Pt(12)
        security_heading.space_after = Pt(8)
        security_section_added = True

        sfr_heading = document.add_paragraph()
        sfr_run = sfr_heading.add_run("5.1 Security Functional Requirements")
        sfr_run.font.size = Pt(18)
        sfr_run.font.bold = True
        sfr_heading.space_before = Pt(8)
        sfr_heading.space_after = Pt(12)

        if payload.sfr_preview_html:
            _append_html_to_document(document, payload.sfr_preview_html)
        else:
            for sfr_item in payload.sfr_list:
                if sfr_item.get('preview'):
                    _append_html_to_document(document, sfr_item['preview'])
                    document.add_paragraph().space_after = Pt(12)

    if payload.sar_preview_html or (payload.sar_list and len(payload.sar_list) > 0):
        if not security_section_added:
            document.add_page_break()
            security_heading = document.add_paragraph()
            security_run = security_heading.add_run("5. Security Requirements")
            security_run.font.size = Pt(20)
            security_run.font.bold = True
            security_heading.space_before = Pt(12)
            security_heading.space_after = Pt(8)
            security_section_added = True

        sar_heading = document.add_paragraph()
        sar_run = sar_heading.add_run("5.2 Security Assurance Requirements")
        sar_run.font.size = Pt(18)
        sar_run.font.bold = True
        sar_heading.space_before = Pt(8)
        sar_heading.space_after = Pt(12)

        if payload.sar_preview_html:
            _append_html_to_document(document, payload.sar_preview_html)
        else:
            if payload.selected_eal:
                eal_para = document.add_paragraph()
                eal_run = eal_para.add_run(f"Evaluation Assurance Level: {payload.selected_eal}")
                eal_run.font.bold = True
                eal_para.space_after = Pt(12)

            for sar_item in payload.sar_list:
                if sar_item.get('preview'):
                    _append_html_to_document(document, sar_item['preview'])
                    document.add_paragraph().space_after = Pt(12)

    if payload.tss_html:
        document.add_page_break()
        tss_heading = document.add_paragraph()
        tss_run = tss_heading.add_run("6. Product Summary Specification")
        tss_run.font.size = Pt(20)
        tss_run.font.bold = True
        tss_heading.space_before = Pt(12)
        tss_heading.space_after = Pt(8)

        intro_paragraph = document.add_paragraph(
            (
                "This section describes the Product security functions that satisfy the security functional requirements. "
                "The Product also includes additional relevant security functions which are also described in the following "
                "sections, as well as a mapping to the security functional requirements satisfied by the Product."
            )
        )
        intro_paragraph.space_after = Pt(12)

        _append_html_to_document(document, payload.tss_html)

    filename = f"{uuid.uuid4().hex}.docx"
    output_path = docx_dir / filename
    document.save(str(output_path))
    return output_path

# CORS configuration: prefer regex if provided to allow any LAN IP on port 5173
origins = os.getenv("CORS_ORIGINS", "http://localhost:5173,http://127.0.0.1:5173").split(",")
origin_regex = os.getenv(
    "CORS_ORIGIN_REGEX",
    None,
)

cors_kwargs = dict(
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

if origin_regex:
    app.add_middleware(
        CORSMiddleware,
        allow_origin_regex=origin_regex,
        **cors_kwargs,
    )
else:
    app.add_middleware(
        CORSMiddleware,
        allow_origins=origins,
        **cors_kwargs,
    )

app.mount("/cover/uploads", StaticFiles(directory=str(COVER_UPLOAD_ROOT)), name="cover-uploads")
app.mount("/cover/docx", StaticFiles(directory=str(COVER_DOCX_ROOT)), name="cover-docx")
app.mount("/security/sfr/docx", StaticFiles(directory=str(SFR_DOCX_ROOT)), name="sfr-docx")
app.mount("/security/sar/docx", StaticFiles(directory=str(SAR_DOCX_ROOT)), name="sar-docx")
app.mount("/spd/docx", StaticFiles(directory=str(SPD_DOCX_ROOT)), name="spd-docx")
app.mount("/so/docx", StaticFiles(directory=str(SO_DOCX_ROOT)), name="so-docx")
app.mount("/tss/docx", StaticFiles(directory=str(TSS_DOCX_ROOT)), name="tss-docx")
app.mount("/st-intro/docx", StaticFiles(directory=str(ST_INTRO_DOCX_ROOT)), name="st-intro-docx")
app.mount("/final-preview/docx", StaticFiles(directory=str(FINAL_DOCX_ROOT)), name="final-docx")


@app.get("/health")
def health(db: Session = Depends(get_db)):
    started = time.time()
    ok = True
    details = {}
    try:
        db.execute(text("SELECT 1"))
    except Exception as e:
        ok = False
        details["db_error"] = str(e)
    latency_ms = int((time.time() - started) * 1000)
    return {
        "status": "ok" if ok else "degraded",
        "latency_ms": latency_ms,
        "database_url": os.getenv("DATABASE_URL", "unset"),
        "timestamp": int(time.time()),
        "details": details,
    }


# CRUD endpoints
@app.get("/components", response_model=List[ComponentOut])
def list_components(
    q: Optional[str] = Query(None, description="Search across select fields"),
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db),
):
    query = db.query(Component)
    if q:
        like = f"%{q}%"
        query = query.filter(
            (
                (Component.class_name.ilike(like))
                | (Component.family.ilike(like))
                | (Component.component.ilike(like))
                | (Component.component_name.ilike(like))
                | (Component.element.ilike(like))
                | (Component.element_item.ilike(like))
            )
        )
    return query.offset(skip).limit(limit).all()


@app.post("/components", response_model=ComponentOut, status_code=201)
def create_component(payload: ComponentCreate, db: Session = Depends(get_db)):
    item = Component(
        class_name=payload.class_name,
        family=payload.family,
        component=payload.component,
        component_name=payload.component_name,
        element=payload.element,
        element_item=payload.element_item,
    )
    db.add(item)
    db.commit()
    db.refresh(item)
    return item


@app.get("/components/{item_id}", response_model=ComponentOut)
def get_component(item_id: int, db: Session = Depends(get_db)):
    item = db.get(Component, item_id)
    if not item:
        raise HTTPException(status_code=404, detail="Not found")
    return item


@app.put("/components/{item_id}", response_model=ComponentOut)
def update_component(item_id: int, payload: ComponentUpdate, db: Session = Depends(get_db)):
    item = db.get(Component, item_id)
    if not item:
        raise HTTPException(status_code=404, detail="Not found")
    data = payload.model_dump(exclude_unset=True, by_alias=True)
    if "class" in data:
        item.class_name = data["class"]
        del data["class"]
    for k, v in data.items():
        setattr(item, k, v)
    db.add(item)
    db.commit()
    db.refresh(item)
    return item


@app.delete("/components/{item_id}", status_code=204)
def delete_component(item_id: int, db: Session = Depends(get_db)):
    item = db.get(Component, item_id)
    if not item:
        raise HTTPException(status_code=404, detail="Not found")
    db.delete(item)
    db.commit()
    return None


@app.post("/cover/upload")
async def upload_cover_image(
    user_id: str = Query(..., alias="user_id"),
    file: UploadFile = File(...),
):
    """Store a temporary cover image for a user session."""

    user_dir = get_user_upload_dir(user_id, create=True)

    allowed_types = {"image/jpeg", "image/png", "image/bmp", "image/x-ms-bmp"}
    if file.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail="Only jpg, jpeg, png, or bmp files are allowed.")

    # Determine file extension based on mime type to avoid trusting filename blindly
    extension_map = {
        "image/jpeg": ".jpg",
        "image/png": ".png",
        "image/bmp": ".bmp",
        "image/x-ms-bmp": ".bmp",
    }
    suffix = extension_map[file.content_type]

    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    # Remove any existing files for the session to keep storage minimal
    if user_dir.exists():
        for existing in user_dir.iterdir():
            if existing.is_file():
                existing.unlink()

    filename = f"{uuid.uuid4().hex}{suffix}"
    destination = user_dir / filename
    with destination.open("wb") as buffer:
        buffer.write(data)

    return {"path": f"/cover/uploads/{user_id}/{filename}"}


@app.post("/cover/preview")
async def generate_cover_preview(payload: CoverPreviewRequest):
    if not payload.user_id:
        raise HTTPException(status_code=400, detail="User identifier is required")

    # Ensure the upload directory exists to maintain parity with the image uploads
    get_user_upload_dir(payload.user_id, create=True)

    output_path = _build_cover_document(payload)
    return {"path": f"/cover/docx/{payload.user_id}/{output_path.name}"}


@app.post("/security/sfr/preview")
async def generate_sfr_preview(payload: HtmlPreviewRequest):
    if not payload.user_id:
        raise HTTPException(status_code=400, detail="User identifier is required")

    output_path = _build_html_preview_document(payload.html_content, payload.user_id, SFR_DOCX_ROOT)
    return {"path": f"/security/sfr/docx/{payload.user_id}/{output_path.name}"}


@app.post("/security/sar/preview")
async def generate_sar_preview(payload: HtmlPreviewRequest):
    if not payload.user_id:
        raise HTTPException(status_code=400, detail="User identifier is required")

    output_path = _build_html_preview_document(payload.html_content, payload.user_id, SAR_DOCX_ROOT)
    return {"path": f"/security/sar/docx/{payload.user_id}/{output_path.name}"}


@app.post("/spd/preview")
async def generate_spd_preview(payload: HtmlPreviewRequest):
    if not payload.user_id:
        raise HTTPException(status_code=400, detail="User identifier is required")

    output_path = _build_html_preview_document(payload.html_content, payload.user_id, SPD_DOCX_ROOT)
    return {"path": f"/spd/docx/{payload.user_id}/{output_path.name}"}


@app.post("/so/preview")
async def generate_security_objectives_preview(payload: HtmlPreviewRequest):
    if not payload.user_id:
        raise HTTPException(status_code=400, detail="User identifier is required")

    output_path = _build_html_preview_document(payload.html_content, payload.user_id, SO_DOCX_ROOT)
    return {"path": f"/so/docx/{payload.user_id}/{output_path.name}"}


@app.post("/tss/preview")
async def generate_tss_preview(payload: HtmlPreviewRequest):
    if not payload.user_id:
        raise HTTPException(status_code=400, detail="User identifier is required")

    output_path = _build_tss_preview_document(payload.html_content, payload.user_id)
    return {"path": f"/tss/docx/{payload.user_id}/{output_path.name}"}


@app.post("/st-intro/preview")
async def generate_st_intro_preview(payload: STIntroPreviewRequest):
    if not payload.user_id:
        raise HTTPException(status_code=400, detail="User identifier is required")

    output_path = _build_st_intro_combined_document(payload)
    return {"path": f"/st-intro/docx/{payload.user_id}/{output_path.name}"}


@app.delete("/cover/upload/{user_id}")
async def cleanup_cover_images(user_id: str):
    """Delete all temporary cover images associated with a user session."""

    user_dir = get_user_upload_dir(user_id, create=False)
    if user_dir.exists():
        shutil.rmtree(user_dir)
    docx_dir = get_user_docx_dir(user_id, create=False)
    if docx_dir.exists():
        shutil.rmtree(docx_dir)
    return {"status": "deleted"}


@app.delete("/cover/preview/{user_id}")
async def cleanup_cover_preview(user_id: str):
    """Delete generated cover preview documents for a user session."""

    docx_dir = get_user_docx_dir(user_id, create=False)
    if docx_dir.exists():
        shutil.rmtree(docx_dir)
    return {"status": "deleted"}


@app.delete("/security/sfr/preview/{user_id}")
async def cleanup_sfr_preview(user_id: str):
    docx_dir = _get_preview_docx_dir(SFR_DOCX_ROOT, user_id, create=False)
    if docx_dir.exists():
        shutil.rmtree(docx_dir)
    return {"status": "deleted"}


@app.delete("/security/sar/preview/{user_id}")
async def cleanup_sar_preview(user_id: str):
    docx_dir = _get_preview_docx_dir(SAR_DOCX_ROOT, user_id, create=False)
    if docx_dir.exists():
        shutil.rmtree(docx_dir)
    return {"status": "deleted"}


@app.delete("/spd/preview/{user_id}")
async def cleanup_spd_preview(user_id: str):
    docx_dir = _get_preview_docx_dir(SPD_DOCX_ROOT, user_id, create=False)
    if docx_dir.exists():
        shutil.rmtree(docx_dir)
    return {"status": "deleted"}


@app.delete("/so/preview/{user_id}")
async def cleanup_security_objectives_preview(user_id: str):
    docx_dir = _get_preview_docx_dir(SO_DOCX_ROOT, user_id, create=False)
    if docx_dir.exists():
        shutil.rmtree(docx_dir)
    return {"status": "deleted"}


@app.delete("/tss/preview/{user_id}")
async def cleanup_tss_preview(user_id: str):
    docx_dir = _get_preview_docx_dir(TSS_DOCX_ROOT, user_id, create=False)
    if docx_dir.exists():
        shutil.rmtree(docx_dir)
    return {"status": "deleted"}


@app.delete("/st-intro/preview/{user_id}")
async def cleanup_st_intro_preview(user_id: str):
    docx_dir = _get_preview_docx_dir(ST_INTRO_DOCX_ROOT, user_id, create=False)
    if docx_dir.exists():
        shutil.rmtree(docx_dir)
    return {"status": "deleted"}


@app.post("/final-preview")
async def generate_final_preview(payload: FinalPreviewRequest):
    if not payload.user_id:
        raise HTTPException(status_code=400, detail="User identifier is required")

    output_path = _build_final_combined_document(payload)
    return {"path": f"/final-preview/docx/{payload.user_id}/{output_path.name}"}


@app.delete("/final-preview/{user_id}")
async def cleanup_final_preview(user_id: str):
    docx_dir = _get_preview_docx_dir(FINAL_DOCX_ROOT, user_id, create=False)
    if docx_dir.exists():
        shutil.rmtree(docx_dir)
    return {"status": "deleted"}


@app.get("/final-preview/download/{user_id}/{filename}")
async def download_final_preview(user_id: str, filename: str):
    if not USER_ID_PATTERN.match(user_id):
        raise HTTPException(status_code=400, detail="Invalid user identifier")

    safe_name = Path(filename).name
    if safe_name != filename:
        raise HTTPException(status_code=400, detail="Invalid file name")

    docx_dir = _get_preview_docx_dir(FINAL_DOCX_ROOT, user_id, create=False)
    file_path = docx_dir / safe_name
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(
        path=str(file_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="Security_Target_Document.docx",
    )
